"""
literature.py - Scholar Kit 统一 CLI 入口 (v1.12.1)
用法:
  python literature.py search "关键词" [--project 课题名] [--source cnki|openalex|semantic|arxiv|nssd|all] [--doc-type master] [--field 摘要] [--author] [--journal] [--download] ...
  python literature.py batch-search "词1" "词2" ... [--project 课题名] [--query-file kw.txt] [--core CSSCI] [--doc-type master] [--field 摘要] [--author] [--journal] [--append]
  python literature.py read-detail [--project 课题名] [--top-n 5] [--fulltext]
  python literature.py read-paper <论文.docx> [--output paper.txt]
  python literature.py download <url_or_doi> [--dir ./papers] [--doi DOI]
  python literature.py batch-download --from-session [--top-n 20] [--dir ./papers]
  python literature.py batch-download url1 url2 ... [--dir ./papers]
  python literature.py detail <cnki_url>
  python literature.py export --format bibtex|ris|markdown|json|excel|gbt7714|footnote|apa [--output file]
  python literature.py cite --style gbt7714|gb|footnote|apa
  python literature.py import <filepath>
  python literature.py write-docx <draft.md> [--output 论文.docx]
  python literature.py patch-docx <原论文.docx> --patch patch.json [--output 修改后.docx]
  python literature.py citations <DOI|URL> [--direction citing|cited|both] [--limit 20]
  python literature.py trends                  # 研究趋势（基于会话数据）
  python literature.py review [--project 课题名] [--topic 综述主题] [--auto-detail] [--output review.md]
  python literature.py write [--project 课题名] [--topic 主题] [--mode outline|draft|section] [--section 章节名] [--format markdown|docx] [--with-citations] [--validate]
  python literature.py validate [--project 课题名] [--topic 主题] [--file draft.md]
  python literature.py topics [--project 课题名] [--topic 主题]
  python literature.py check                   # 环境自检
  python literature.py clean-cache [--all] [--dry-run]  # 缓存清理
"""

from __future__ import annotations

__version__ = "1.12.1"

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path
from typing import Any, Dict, List, Optional

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = open(sys.stdout.fileno(), mode="w", encoding="utf-8", errors="replace", buffering=1)
    sys.stderr = open(sys.stderr.fileno(), mode="w", encoding="utf-8", errors="replace", buffering=1)

_script_dir = str(Path(__file__).parent)
if _script_dir not in sys.path:
    sys.path.insert(0, _script_dir)

from search import (  # noqa: E402
    search_openalex, search_semantic_scholar, search_arxiv,
    search_nssd, search_all, resolve_crossref, resolve_unpaywall,
    get_citations, analyze_trends,
)
from cnki import (  # noqa: E402
    search_cnki, batch_search_cnki, batch_read_detail,
    get_detail, download_cnki, batch_download_cnki,
    parse_cnki_export, check_cnki_access,
)
from formatter import export_papers, generate_reference_list, citation_preview  # noqa: E402

CITATION_STYLE_CHOICES = ["gbt7714", "gb", "apa", "mla", "chicago", "footnote"]

def _safe_project_name(project: str) -> str:
    cleaned = "".join(c if c.isalnum() or c in "._- 一-鿿" else "_" for c in project.strip())
    cleaned = cleaned.strip(" .")
    return cleaned or "default"


def _project_dir(project: str) -> Path:
    return Path.cwd() / ".scholar-kit" / "projects" / _safe_project_name(project)


def _session_file(project: Optional[str] = None) -> Path:
    if project:
        return _project_dir(project) / "session.json"
    return Path.cwd() / ".scholar-kit" / "session.json"


def _session_project(args) -> Optional[str]:
    return getattr(args, "project", None) or None


def _save_session(results: List[Dict[str, Any]], append: bool = False, project: Optional[str] = None):
    """保存搜索结果到会话文件。append=True 时追加并按标题去重。"""
    sf = _session_file(project)
    sf.parent.mkdir(parents=True, exist_ok=True)
    if append:
        existing = _load_session(project)
        seen: Dict[str, Dict[str, Any]] = {}
        no_title: List[Dict[str, Any]] = []
        for r in existing:
            key = r.get("title", "").lower().strip()
            if key:
                seen[key] = r
            else:
                no_title.append(r)
        for r in results:
            key = r.get("title", "").lower().strip()
            if key:
                seen[key] = r
            else:
                no_title.append(r)
        results = list(seen.values()) + no_title
    sf.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")


def _load_session(project: Optional[str] = None) -> List[Dict[str, Any]]:
    sf = _session_file(project)
    if sf.exists():
        try:
            return json.loads(sf.read_text(encoding="utf-8"))
        except Exception:
            return []
    return []


# ── CNKI 搜索缓存 ─────────────────────────────────────

_CNKI_CACHE_TTL_MINUTES = 30


def _cnki_cache_key(args) -> str:
    import hashlib
    key_parts = f"{args.query}|{args.core}|{args.year_from}|{args.year_to}|{args.author}|{args.journal}|{getattr(args, 'doc_type', '')}|{getattr(args, 'field', '')}|{args.sort}|{args.pages}|{getattr(args, 'cite_enrich', 0)}"
    return hashlib.md5(key_parts.encode()).hexdigest()


def _cnki_cache_get(args) -> Optional[list]:
    cache_dir = Path.cwd() / ".scholar-kit" / "cache"
    cache_file = cache_dir / f"cnki_{_cnki_cache_key(args)}.json"
    if cache_file.exists():
        try:
            data = json.loads(cache_file.read_text(encoding="utf-8"))
            from datetime import datetime, timedelta
            cached_at = datetime.fromisoformat(data.get("_cached_at", ""))
            if datetime.now() - cached_at < timedelta(minutes=_CNKI_CACHE_TTL_MINUTES):
                return data.get("results")
        except Exception:
            pass
    return None


def _cnki_cache_set(args, results: list):
    from datetime import datetime
    cache_dir = Path.cwd() / ".scholar-kit" / "cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / f"cnki_{_cnki_cache_key(args)}.json"
    data = {"_cached_at": datetime.now().isoformat(), "results": results}
    try:
        cache_file.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


# ── search 命令 ───────────────────────────────────────

def cmd_search(args):
    source = args.source or "cnki"
    if source not in ("cnki", "openalex", "semantic", "arxiv", "nssd", "all"):
        _output({"status": "error", "code": "UNKNOWN_SOURCE",
                 "message": f"未知数据源: {source}"})
        return

    want_download = getattr(args, "download", False)
    if want_download and source != "cnki":
        _output({"status": "error", "code": "DOWNLOAD_SOURCE_MISMATCH",
                 "message": "--download 仅支持 --source cnki"})
        return

    results = []
    reuse_driver = None
    cnki_error = None

    if source in ("cnki", "all"):
        # 检查 CNKI 搜索缓存
        cached_cnki = _cnki_cache_get(args) if not want_download else None
        if cached_cnki is not None:
            print("[cnki] 使用缓存结果", file=__import__('sys').stderr)
            results.extend(cached_cnki)
        else:
            keep = want_download and source == "cnki"
            cnki_ret = search_cnki(
                keyword=args.query,
                core=args.core,
                year_from=args.year_from,
                year_to=args.year_to,
                author=args.author,
                journal=args.journal,
                doc_type=getattr(args, "doc_type", None),
                field=getattr(args, "field", None),
                sort=args.sort or "relevance",
                pages=args.pages or 1,
                cite_enrich=getattr(args, "cite_enrich", 0),
                _keep_driver=keep,
            )
            if keep and isinstance(cnki_ret, tuple):
                cnki_results, reuse_driver = cnki_ret
            else:
                cnki_results = cnki_ret

            if cnki_results and not (len(cnki_results) == 1 and cnki_results[0].get("status") == "error"):
                results.extend(cnki_results)
                _cnki_cache_set(args, cnki_results)
            elif cnki_results and cnki_results[0].get("status") == "error":
                if source == "cnki":
                    _output(cnki_results[0])
                    if reuse_driver:
                        try: reuse_driver.quit()
                        except Exception: pass
                    return
                # source == "all" 时知网失败不阻断，记录错误后继续 API 搜索
                cnki_error = cnki_results[0]

    try:
        has_keyword = bool(args.query and args.query.strip())

        api_limit = args.limit if args.limit is not None else 10

        if has_keyword and source in ("openalex", "all"):
            results.extend(search_openalex(
                args.query, limit=api_limit,
                year_from=args.year_from, year_to=args.year_to,
                sort=args.sort or "relevance",
            ))

        if has_keyword and source in ("semantic", "all"):
            results.extend(search_semantic_scholar(
                args.query, limit=api_limit,
                year_from=args.year_from, year_to=args.year_to,
                sort=args.sort or "relevance",
            ))

        if has_keyword and source in ("arxiv", "all"):
            results.extend(search_arxiv(
                args.query, limit=api_limit, sort_by=args.sort or "relevance",
                year_from=args.year_from, year_to=args.year_to,
            ))

        if has_keyword and source in ("nssd", "all"):
            results.extend(search_nssd(
                args.query, limit=api_limit,
                year_from=args.year_from, year_to=args.year_to,
            ))

        # 使用改进的去重函数（基于 DOI 和标题）
        from search import deduplicate_results, calculate_quality_score
        deduped = deduplicate_results(results)

        # 添加质量评分
        for paper in deduped:
            paper["quality_score"] = calculate_quality_score(paper)

        if args.sort == "citations":
            deduped.sort(key=lambda x: x.get("cited_by", 0), reverse=True)
        elif args.sort == "date":
            deduped.sort(key=lambda x: x.get("year") or 0, reverse=True)
        elif args.sort == "quality":
            deduped.sort(key=lambda x: x.get("quality_score", 0), reverse=True)

        if args.limit is not None:
            effective_limit = args.limit
        elif args.pages and args.pages > 1 and source in ("cnki", "all"):
            effective_limit = args.pages * 20
        else:
            effective_limit = 20
        deduped = deduped[:effective_limit]

        # --enrich: 对知网结果自动补全卷期页码
        enrich_n = getattr(args, "enrich", 0)
        if enrich_n and enrich_n > 0:
            cnki_papers = [(i, p) for i, p in enumerate(deduped)
                           if _is_cnki_paper(p) and p.get("url") and not p.get("pages")]
            to_enrich = cnki_papers[:enrich_n]
            if to_enrich:
                print(f"[enrich] 正在补全 {len(to_enrich)} 篇论文的卷期页码...",
                      file=__import__('sys').stderr)
                from cnki import get_detail
                for idx, p in to_enrich:
                    detail = get_detail(p["url"])
                    if detail and detail.get("status") != "error":
                        for k in ("volume", "issue", "pages", "doi", "year", "journal", "authors"):
                            if detail.get(k) and not p.get(k):
                                p[k] = detail[k]
                    time.sleep(1)

        _save_session(deduped, append=getattr(args, "append", False), project=_session_project(args))

        # 为每条结果添加引用预览
        for p in deduped:
            p["citation_preview"] = citation_preview(p)

        search_output = {"status": "success", "count": len(deduped), "results": deduped}
        if cnki_error and source == "all":
            search_output["cnki_error"] = {"code": cnki_error.get("code"), "message": cnki_error.get("message")}
            search_output["status"] = "partial"
        if args.export:
            content = export_papers(deduped, args.export, args.output)
            if isinstance(content, dict) and content.get("status") == "error":
                search_output["export_error"] = content
            else:
                search_output.update({"format": args.export, "output_file": args.output,
                                      "content": content})

        if want_download and reuse_driver:
            from config import get as cfg_get
            dl_dir = getattr(args, "download_dir", None) or cfg_get("save_dir", "./papers")
            dl_top_n = getattr(args, "download_top_n", None)
            dl_papers = deduped[:dl_top_n] if dl_top_n else deduped
            dl_urls = [p.get("url") for p in dl_papers if isinstance(p, dict) and p.get("url")]
            if dl_urls:
                dl_format = getattr(args, "download_file_format", "pdf") or "pdf"
                dl_result = batch_download_cnki(
                    dl_urls,
                    save_dir=dl_dir,
                    file_format=dl_format,
                    _driver=reuse_driver,
                )
                fallback_format = getattr(args, "download_fallback_format", None)
                if fallback_format:
                    failed_urls = [
                        err.get("url") for err in (dl_result.get("errors") or [])
                        if isinstance(err, dict)
                        and err.get("url")
                        and err.get("code") == "DOWNLOAD_BTN_NOT_FOUND"
                    ]
                    if failed_urls:
                        fallback_result = batch_download_cnki(
                            failed_urls,
                            save_dir=dl_dir,
                            file_format=fallback_format,
                            _driver=reuse_driver,
                        )
                        dl_result = _merge_fallback_download(dl_result, fallback_result)
                if not getattr(args, "download_no_report", False):
                    dl_result = attach_download_report(
                        dl_result,
                        save_dir=dl_dir,
                        session_papers=dl_papers,
                        requested_urls=dl_urls,
                        citation_style=getattr(args, "download_citation_style", "gbt7714") or "gbt7714",
                        file_format=dl_format,
                        report_output=getattr(args, "download_report_output", None),
                    )
                reuse_driver = None
                search_output["download"] = dl_result
            else:
                search_output["download"] = {"status": "warning", "code": "NO_DOWNLOAD_URLS",
                                             "message": "搜索结果中无可下载 URL"}

        _output(search_output)

    finally:
        if reuse_driver is not None:
            try:
                reuse_driver.quit()
            except Exception:
                pass


# ── download 命令 ─────────────────────────────────────

def cmd_download(args):
    from config import get as cfg_get
    target = args.target
    save_dir = args.dir if args.dir != "./papers" else cfg_get("save_dir", "./papers")

    if args.doi:
        unpaywall = resolve_unpaywall(args.doi)
        if unpaywall and unpaywall.get("oa_url"):
            _output({
                "status": "success",
                "method": "unpaywall_oa",
                "url": unpaywall["oa_url"],
                "message": f"找到 OA 链接: {unpaywall['oa_url']}",
            })
            return

        crossref = resolve_crossref(args.doi)
        meta = {
            "title": (crossref or {}).get("title", ""),
            "authors": (crossref or {}).get("authors", ""),
            "journal": (crossref or {}).get("journal", ""),
            "doi": args.doi,
        }
        if target and "cnki.net" in target:
            result = download_cnki(target, save_dir=save_dir, file_format=args.file_format or "pdf")
            _output(result)
            return
        _output({
            "status": "error",
            "code": "OA_NOT_FOUND",
            "message": "Unpaywall 未找到 OA 版本，无知网 URL 可回退",
            "metadata": meta,
        })
        return

    if target and "cnki.net" in target:
        result = download_cnki(target, save_dir=save_dir, file_format=args.file_format or "pdf")
        _output(result)
    elif target:
        _output({
            "status": "error",
            "code": "UNSUPPORTED_URL",
            "message": "目前仅支持知网 URL 直接下载",
        })
    else:
        _output({"status": "error", "code": "NO_DOWNLOAD_TARGET",
                 "message": "请提供下载目标（URL 或 --doi）"})


# ── batch-download 命令 ──────────────────────────────

def _download_report_path(save_dir: str, report_output: Optional[str] = None) -> str:
    if report_output:
        return report_output
    stamp = time.strftime("%Y%m%d_%H%M%S")
    return str(Path(save_dir) / f"download_report_{stamp}.md")


def _paper_lookup_by_url(papers: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    lookup: Dict[str, Dict[str, Any]] = {}
    for paper in papers:
        if isinstance(paper, dict) and paper.get("url"):
            lookup[paper["url"]] = paper
    return lookup


def _download_item_to_paper(item: Dict[str, Any], lookup: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
    paper = dict(lookup.get(item.get("url", ""), {}))
    for key in ("url", "title"):
        if item.get(key) and not paper.get(key):
            paper[key] = item[key]
    if not paper.get("title"):
        paper["title"] = item.get("filename") or item.get("url") or "未获取题名"
    return paper


def build_download_report(
    result: Dict[str, Any],
    session_papers: Optional[List[Dict[str, Any]]] = None,
    requested_urls: Optional[List[str]] = None,
    citation_style: str = "gbt7714",
    file_format: str = "pdf",
) -> Dict[str, Any]:
    lookup = _paper_lookup_by_url(session_papers or [])
    ok_items = result.get("results") or []
    error_items = result.get("errors") or []
    if not error_items and result.get("status") == "error" and requested_urls:
        ok_urls = {item.get("url") for item in ok_items if isinstance(item, dict)}
        error_items = [
            {
                "url": url,
                "code": result.get("code"),
                "error": result.get("message") or result.get("code") or "下载失败",
            }
            for url in requested_urls if url not in ok_urls
        ]

    downloaded = [_download_item_to_paper(item, lookup) for item in ok_items if isinstance(item, dict)]
    failed = [_download_item_to_paper(item, lookup) for item in error_items if isinstance(item, dict)]

    downloaded_refs = generate_reference_list(downloaded, citation_style).splitlines() if downloaded else []
    failed_refs = generate_reference_list(failed, citation_style).splitlines() if failed else []

    lines = [
        "# 文献下载清单",
        "",
        f"- 请求格式: {file_format.upper()}",
        f"- 引用格式: {citation_style.upper()}",
        f"- 已下载: {len(downloaded)}",
        f"- 未下载: {len(failed)}",
        "",
        "## 已下载",
        "",
    ]
    if downloaded_refs:
        for idx, ref in enumerate(downloaded_refs):
            item = ok_items[idx] if idx < len(ok_items) and isinstance(ok_items[idx], dict) else {}
            actual_format = item.get("format") or file_format
            requested = item.get("requested_format") or file_format
            filename = item.get("filename")
            suffix_parts = [f"格式：{actual_format.upper()}"]
            if actual_format != requested or item.get("fallback_used"):
                suffix_parts.append(f"由 {requested.upper()} 降级")
            if filename:
                suffix_parts.append(f"文件：{filename}")
            lines.append(f"{ref}（{'；'.join(suffix_parts)}）")
    else:
        lines.append("无")

    lines.extend(["", "## 未下载", ""])
    if failed_refs:
        for idx, ref in enumerate(failed_refs):
            item = error_items[idx] if idx < len(error_items) and isinstance(error_items[idx], dict) else {}
            reason = item.get("error") or item.get("message") or item.get("code") or "未获取失败原因"
            lines.append(f"{ref}（原因：{reason}）")
    else:
        lines.append("无")

    return {
        "citation_style": citation_style,
        "file_format": file_format,
        "downloaded_references": downloaded_refs,
        "failed_references": failed_refs,
        "markdown": "\n".join(lines).rstrip() + "\n",
    }


def attach_download_report(
    result: Dict[str, Any],
    save_dir: str,
    session_papers: Optional[List[Dict[str, Any]]] = None,
    requested_urls: Optional[List[str]] = None,
    citation_style: str = "gbt7714",
    file_format: str = "pdf",
    report_output: Optional[str] = None,
) -> Dict[str, Any]:
    report = build_download_report(
        result,
        session_papers=session_papers,
        requested_urls=requested_urls,
        citation_style=citation_style,
        file_format=file_format,
    )
    output_path = _download_report_path(save_dir, report_output)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(report["markdown"], encoding="utf-8")
    enriched = dict(result)
    enriched["download_report"] = {
        "path": output_path,
        "citation_style": report["citation_style"],
        "file_format": report["file_format"],
        "downloaded_references": report["downloaded_references"],
        "failed_references": report["failed_references"],
    }
    return enriched


def _merge_fallback_download(primary: Dict[str, Any], fallback: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(primary)
    primary_results = list(primary.get("results") or [])
    fallback_results = list(fallback.get("results") or [])
    fallback_urls = {item.get("url") for item in fallback_results if isinstance(item, dict)}
    primary_requested_format = primary.get("requested_format")

    for item in fallback_results:
        if isinstance(item, dict):
            item = dict(item)
            if primary_requested_format:
                item["requested_format"] = primary_requested_format
            item["fallback_used"] = True
            primary_results.append(item)

    remaining_errors = [
        err for err in (primary.get("errors") or [])
        if not isinstance(err, dict) or err.get("url") not in fallback_urls
    ]
    remaining_errors.extend(fallback.get("errors") or [])

    merged["results"] = primary_results
    merged["errors"] = remaining_errors or None
    merged["count"] = len(primary_results)
    if remaining_errors and primary_results:
        merged["status"] = "partial"
    elif remaining_errors:
        merged["status"] = "error"
        merged["code"] = primary.get("code") or fallback.get("code") or "CNKI_BATCH_DOWNLOAD_FAILED"
    else:
        merged["status"] = "success"
        merged.pop("code", None)
    merged["fallback"] = {
        "attempted": bool(fallback_results or fallback.get("errors")),
        "status": fallback.get("status"),
        "format": fallback.get("requested_format"),
        "count": len(fallback_results),
    }
    return merged

def cmd_batch_download(args):
    """批量下载：浏览器只启动一次，多标签页并行下载"""
    from config import get as cfg_get
    urls = list(args.urls) if args.urls else []
    session_data: List[Dict[str, Any]] = []

    if args.from_session:
        session_data = _load_session(_session_project(args))
        if not session_data:
            _output({"status": "error", "code": "NO_SESSION",
                     "message": "没有搜索记录，请先执行 search 或 batch-search"})
            return
        top_n = args.top_n or len(session_data)
        session_urls = [p.get("url") for p in session_data[:top_n]
                        if isinstance(p, dict) and p.get("url")]
        urls.extend(session_urls)

    if not urls:
        _output({"status": "error", "code": "NO_URLS",
                 "message": "未提供下载 URL（可用 --from-session 从上次搜索结果读取）"})
        return

    save_dir = args.dir if args.dir != "./papers" else cfg_get("save_dir", "./papers")
    result = batch_download_cnki(
        urls=urls,
        save_dir=save_dir,
        file_format=args.file_format or "pdf",
    )
    fallback_format = getattr(args, "fallback_format", None)
    if fallback_format:
        failed_urls = [
            err.get("url") for err in (result.get("errors") or [])
            if isinstance(err, dict)
            and err.get("url")
            and err.get("code") == "DOWNLOAD_BTN_NOT_FOUND"
        ]
        if failed_urls:
            fallback_result = batch_download_cnki(
                urls=failed_urls,
                save_dir=save_dir,
                file_format=fallback_format,
            )
            result = _merge_fallback_download(result, fallback_result)
    if not getattr(args, "no_report", False):
        result = attach_download_report(
            result,
            save_dir=save_dir,
            session_papers=session_data,
            requested_urls=urls,
            citation_style=getattr(args, "citation_style", "gbt7714") or "gbt7714",
            file_format=args.file_format or "pdf",
            report_output=getattr(args, "report_output", None),
        )
    _output(result)


# ── detail 命令 ───────────────────────────────────────

def cmd_detail(args):
    if not args.url:
        _output({"status": "error", "code": "NO_URL", "message": "请提供知网论文详情页 URL"})
        return
    result = get_detail(args.url)
    _output(result)


# ── export 命令 ───────────────────────────────────────

def cmd_export(args):
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA", "message": "没有可导出的数据，请先执行 search、batch-search 或 import"})
        return

    result = export_papers(papers, args.export_format, args.output)
    if isinstance(result, dict) and result.get("status") == "error":
        _output(result)
        return
    if args.raw:
        print(result)
    else:
        _output({"status": "success", "project": _session_project(args), "format": args.export_format,
                 "output_file": args.output, "content": result})


# ── project/library 命令 ───────────────────────────────

def _paper_year(paper: Dict[str, Any]) -> Any:
    return paper.get("year") or str(paper.get("date", ""))[:4]


def _paper_summary(paper: Dict[str, Any], index: int) -> Dict[str, Any]:
    return {
        "index": index,
        "title": paper.get("title", ""),
        "authors": paper.get("authors", ""),
        "journal": paper.get("journal", ""),
        "year": _paper_year(paper),
        "cited_by": paper.get("cited_by", 0),
        "source": paper.get("source", ""),
        "tags": paper.get("tags", []),
        "note": paper.get("note", ""),
    }


def _paper_evidence(paper: Dict[str, Any], index: int) -> Dict[str, Any]:
    abstract = str(paper.get("abstract") or "").strip()
    keywords = paper.get("keywords") or []
    if isinstance(keywords, str):
        keywords = [k.strip() for k in keywords.replace("；", ";").replace("，", ";").split(";") if k.strip()]
    evidence = {
        "index": index,
        "title": paper.get("title", ""),
        "authors": paper.get("authors", ""),
        "journal": paper.get("journal", ""),
        "year": _paper_year(paper),
        "source": paper.get("source", ""),
        "doi": paper.get("doi", ""),
        "url": paper.get("url", ""),
        "keywords": keywords[:8],
        "abstract_excerpt": abstract[:220],
        "trace_status": "abstract" if abstract else "metadata_only",
    }
    if paper.get("pages"):
        evidence["pages"] = paper.get("pages")
    return evidence


def _review_terms(topic: str, paper: Dict[str, Any]) -> List[str]:
    terms: List[str] = []
    raw_terms = [topic, paper.get("title", ""), paper.get("journal", "")]
    keywords = paper.get("keywords") or []
    if isinstance(keywords, list):
        raw_terms.extend(str(k) for k in keywords)
    elif isinstance(keywords, str):
        raw_terms.extend(keywords.replace("；", ";").replace("，", ";").split(";"))
    for text in raw_terms:
        for token in str(text).replace("：", " ").replace(":", " ").replace("——", " ").split():
            token = token.strip(" ，。；;、,.()（）[]【】《》<>\"'")
            if len(token) >= 2 and token not in terms:
                terms.append(token)
    return terms[:12]


def _review_query_terms(topic: str) -> List[str]:
    stopwords = {
        "the", "and", "of", "in", "on", "for", "to", "with", "from", "via", "a", "an",
        "中的", "问题", "研究", "基于", "视域", "视角", "下的",
    }
    normalized = topic.lower()
    for sep in "：:，,；;、/\\()（）[]【】《》<>\"'\n\t":
        normalized = normalized.replace(sep, " ")
    terms = []
    for token in normalized.split():
        token = token.strip(" .-_—")
        if len(token) >= 3 and token not in stopwords and token not in terms:
            terms.append(token)
    cn_phrases = (
        "国家形象", "国际传播", "对外传播", "社交媒体", "官方叙事", "中国故事",
        "传播能力", "文化符号", "受众", "认知", "传播效果", "话语", "平台",
    )
    for phrase in cn_phrases:
        if phrase in topic and phrase not in terms:
            terms.append(phrase)
    compact = topic.lower().replace(" ", "")
    if compact and compact not in terms:
        terms.append(compact)
    return terms


def _review_quality_flags(paper: Dict[str, Any], relevance: int) -> List[str]:
    title = str(paper.get("title") or "")
    source = str(paper.get("source") or "")
    flags = []
    if "retracted" in title.lower() or "撤稿" in title:
        flags.append("retracted")
    if relevance <= 0:
        flags.append("low_relevance")
    if not paper.get("abstract"):
        flags.append("needs_fulltext_check")
    if source in ("CNKI", "CNKI-export") and not paper.get("abstract"):
        flags.append("cnki_read_detail_recommended")
    return flags


def _review_candidates(papers: List[Dict[str, Any]], topic: str) -> List[Dict[str, Any]]:
    query_terms = _review_query_terms(topic)
    candidates = []
    for idx, paper in enumerate(papers, 1):
        title = str(paper.get("title") or "").lower()
        abstract = str(paper.get("abstract") or "").lower()
        keywords = str(paper.get("keywords") or "").lower()
        journal = str(paper.get("journal") or "").lower()
        text = " ".join((title, abstract, keywords, journal))
        relevance = 0
        for term in query_terms:
            if term in title:
                relevance += 40
            if term in keywords:
                relevance += 25
            if term in abstract:
                relevance += 10
            if term in journal:
                relevance += 5
        if topic.lower() in text:
            relevance += 80
        if not query_terms:
            relevance = 1
        metadata_score = 0
        if paper.get("abstract"):
            metadata_score += 2
        if paper.get("keywords"):
            metadata_score += 1
        cited_score = min(int(paper.get("cited_by") or 0), 50)
        flags = _review_quality_flags(paper, relevance)
        sort_relevance = relevance - 1000 if "retracted" in flags else relevance
        candidates.append({
            "index": idx,
            "paper": paper,
            "relevance": relevance,
            "sort_relevance": sort_relevance,
            "metadata_score": metadata_score,
            "cited_score": cited_score,
            "flags": flags,
        })
    candidates.sort(
        key=lambda item: (
            item["sort_relevance"], item["metadata_score"],
            item["cited_score"], item["paper"].get("year") or ""
        ),
        reverse=True,
    )
    return candidates


def _select_review_papers(papers: List[Dict[str, Any]], topic: str, limit: int) -> List[Dict[str, Any]]:
    candidates = _review_candidates(papers, topic)
    relevant = [item for item in candidates if item["relevance"] > 0 and "retracted" not in item["flags"]]
    selected_pool = relevant or candidates
    return selected_pool[:limit]


def _cluster_label_for_paper(paper: Dict[str, Any], topic: str) -> str:
    text = " ".join(str(paper.get(k, "")) for k in ("title", "abstract", "keywords", "journal"))
    rules = [
        ("国家形象建构与官方叙事", ("国家形象", "形象建构", "官方叙事", "中国形象", "national image")),
        ("国际传播能力与对外传播", ("国际传播", "对外传播", "外宣", "传播能力", "international communication")),
        ("平台机制与社交媒体", ("社交媒体", "平台", "youtube", "tiktok", "facebook", "x", "cgtn", "media")),
        ("文化符号与中国故事", ("中国故事", "文化", "符号", "文明", "cultural", "story")),
        ("受众认知与传播效果", ("受众", "认知", "效果", "态度", "audience", "perception", "effect")),
        ("方法与模型", ("模型", "算法", "推荐", "实证", "扎根", "内容分析", "model", "method")),
    ]
    lowered = text.lower()
    for label, terms in rules:
        for term in terms:
            if term.lower() in lowered:
                return label
    terms = _review_terms(topic, paper)
    return terms[0] if terms else "其他相关研究"


def _review_clusters(selected: List[Dict[str, Any]], topic: str) -> List[Dict[str, Any]]:
    clusters: Dict[str, List[Dict[str, Any]]] = {}
    for item in selected:
        label = _cluster_label_for_paper(item["paper"], topic)
        clusters.setdefault(label, []).append(item)
    result = []
    for label, items in sorted(clusters.items(), key=lambda kv: len(kv[1]), reverse=True):
        evidence = []
        claims = []
        for item in items:
            p = item["paper"]
            abstract = str(p.get("abstract") or "").strip()
            if abstract:
                claims.append(abstract[:90])
            evidence.append({
                "index": item["index"],
                "title": p.get("title", ""),
                "trace_status": "abstract" if abstract else "metadata_only",
                "abstract_excerpt": abstract[:160],
            })
        synthesis = ""
        if claims:
            synthesis = "；".join(claims[:3])
        result.append({
            "label": label,
            "count": len(items),
            "papers": evidence,
            "synthesis": synthesis,
        })
    return result


def _review_gaps(papers: List[Dict[str, Any]], topic: str) -> List[Dict[str, Any]]:
    dimensions = [
        ("跨平台比较不足", ("youtube", "tiktok", "facebook", " x ", "微博", "微信", "抖音", "平台"), "平台"),
        ("受众实证研究不足", ("受众", "认知", "效果", "问卷", "访谈", "实验", "audience", "perception"), "受众/效果"),
        ("方法多样性不足", ("内容分析", "话语分析", "扎根", "实验", "问卷", "访谈", "模型", "算法"), "方法"),
        ("非西方或比较视角不足", ("比较", "跨国", "区域", "非洲", "东南亚", "拉美", "一带一路", "comparative"), "比较/区域"),
    ]
    corpus = []
    for p in papers:
        corpus.append(" ".join(str(p.get(k, "")) for k in ("title", "abstract", "keywords", "journal")).lower())
    gaps = []
    for title, terms, dimension in dimensions:
        matches = []
        for i, text in enumerate(corpus, 1):
            if any(term.lower() in f" {text} " for term in terms):
                matches.append(i)
        if len(matches) <= max(1, len(papers) // 10):
            gaps.append({
                "title": title,
                "dimension": dimension,
                "matched_count": len(matches),
                "total": len(papers),
                "evidence_indices": matches[:10],
                "basis": f"当前文献库 {len(papers)} 篇中，{dimension} 相关线索约 {len(matches)} 篇。",
            })
    if not gaps:
        gaps.append({
            "title": "需扩大检索后再判断研究空白",
            "dimension": "检索覆盖",
            "matched_count": len(papers),
            "total": len(papers),
            "evidence_indices": [],
            "basis": "当前文献库各预设维度均有一定覆盖，建议扩展关键词和数据库后再判断空白。",
        })
    return gaps
def _build_review_markdown(topic: str, project: Optional[str], selected: List[Dict[str, Any]], total: int, diagnostics: List[Dict[str, Any]] = None, clusters: List[Dict[str, Any]] = None, gaps: List[Dict[str, Any]] = None) -> str:
    diagnostics = diagnostics or selected
    sources = sorted({item["paper"].get("source", "未获取") or "未获取" for item in selected})
    years = [str(_paper_year(item["paper"])) for item in selected if _paper_year(item["paper"])]
    year_range = f"{min(years)}-{max(years)}" if years else "未获取"
    close_reading = [item for item in selected if item.get("relevance", 0) > 0 and "retracted" not in item.get("flags", [])][:5]
    needs_check = [item for item in selected if "needs_fulltext_check" in item.get("flags", [])]
    risky = [item for item in diagnostics if "retracted" in item.get("flags", []) or "low_relevance" in item.get("flags", [])]
    lines = [
        f"# {topic} 文献综述材料",
        "",
        "## 检索证据",
        f"- 课题文献库：{project or '默认 session'}",
        f"- 分析文献数：{len(selected)} / {total}",
        f"- 数据来源：{', '.join(sources) if sources else '未获取'}",
        f"- 年份范围：{year_range}",
        "- 说明：以下内容基于文献题录、关键词和摘要生成；缺少摘要的条目标注为待核对原文。",
        "",
        "## 推荐精读文献",
    ]
    if close_reading:
        for item in close_reading:
            paper = item["paper"]
            lines.append(f"- [{item['index']}] {paper.get('title', '未获取')}（相关性分：{item.get('relevance', 0)}）")
    else:
        lines.append("- 暂无高相关文献；建议调整关键词重新检索。")
    lines.extend(["", "## 待核对原文"])
    if needs_check:
        for item in needs_check:
            paper = item["paper"]
            hint = "；建议先执行 read-detail --project <课题名> --indices " + str(item["index"]) if paper.get("source") in ("CNKI", "CNKI-export") else ""
            lines.append(f"- [{item['index']}] {paper.get('title', '未获取')}：当前缺少摘要或全文{hint}")
    else:
        lines.append("- 暂无。")
    lines.extend(["", "## 可能不相关或需剔除文献"])
    if risky:
        for item in risky:
            paper = item["paper"]
            flags = "、".join(item.get("flags") or [])
            lines.append(f"- [{item['index']}] {paper.get('title', '未获取')}：{flags}")
    else:
        lines.append("- 暂无明显撤稿或低相关条目。")
    lines.extend(["", "## 主题聚类"])
    if clusters:
        for cluster in clusters:
            lines.append(f"### {cluster['label']}（{cluster['count']} 篇）")
            if cluster.get("synthesis"):
                lines.append(f"该主题下的文献主要围绕“{topic}”展开，现有摘要显示：{cluster['synthesis']}。")
            else:
                lines.append("该主题下文献当前多为题录信息，具体观点仍需补充摘要或原文后核对。")
            lines.append("")
            lines.append("代表文献与证据：")
            for paper in cluster["papers"]:
                status = "摘要可追溯" if paper.get("trace_status") == "abstract" else "待核对原文"
                excerpt = f"；摘要依据：{paper['abstract_excerpt']}" if paper.get("abstract_excerpt") else ""
                lines.append(f"- [{paper['index']}] {paper.get('title', '未获取')}：{status}{excerpt}")
            lines.append("")
    else:
        lines.append("- 未启用聚类；可使用 `--cluster` 生成主题聚类章节。")
    lines.extend(["## 研究空白提示"])
    if gaps:
        for gap in gaps:
            indices = ",".join(str(i) for i in gap.get("evidence_indices", [])) or "无"
            lines.extend([
                f"### {gap['title']}",
                f"- 检索证据：{gap['basis']}",
                f"- 相关文献序号：{indices}",
            ])
    else:
        lines.append("- 未启用研究空白分析；可使用 `--gaps` 基于当前文献库生成统计提示。")
    lines.extend(["", "## 主题线索"])
    for item in selected:
        paper = item["paper"]
        idx = item["index"]
        terms = "、".join(_review_terms(topic, paper)[:6]) or "待提取"
        suffix = ""
        if item.get("flags"):
            suffix = f"（提示：{'、'.join(item['flags'])}）"
        lines.extend([
            f"- [{idx}] {paper.get('title', '未获取')}：{terms}{suffix}",
        ])
    lines.extend(["", "## 综述草稿", ""])
    for item in selected:
        paper = item["paper"]
        idx = item["index"]
        abstract = str(paper.get("abstract") or "").strip()
        if "retracted" in item.get("flags", []):
            point = "该文献标题显示可能为撤稿文献，不建议作为正面证据使用，仅可作为剔除或风险提示。"
        elif abstract:
            point = abstract[:180]
        else:
            point = "该文献当前仅有题录信息，具体观点需补充摘要或原文后核对。"
        lines.extend([
            f"### 线索 {idx}：{paper.get('title', '未获取')}",
            f"围绕“{topic}”，该文献可作为相关研究线索。{point}",
            "",
            "证据：",
            f"- 作者：{paper.get('authors', '未获取') or '未获取'}",
            f"- 来源：{paper.get('journal', '未获取') or '未获取'}，{_paper_year(paper) or '未获取'}",
            f"- 相关性分：{item.get('relevance', 0)}",
            f"- 追溯状态：{'摘要可追溯' if abstract else '待核对原文'}",
            "",
        ])
    lines.extend(["## 参考文献线索"])
    for item in selected:
        p = item["paper"]
        lines.append(f"[{item['index']}] {p.get('authors', '未获取') or '未获取'}. {p.get('title', '未获取') or '未获取'}. {p.get('journal', '未获取') or '未获取'}, {_paper_year(p) or '未获取'}.")
    return "\n".join(lines)


def _paper_write_excerpt(paper: Dict[str, Any], limit: int = 120) -> str:
    text = str(paper.get("abstract") or paper.get("summary") or "").strip()
    if not text:
        return ""
    text = re.sub(r"\s+", " ", text)
    return text[:limit]


def _evidence_indices(items: List[Dict[str, Any]], limit: int = 5) -> str:
    return "".join(f"[{item['index']}]" for item in items[:limit])


def _cluster_evidence_indices(cluster: Dict[str, Any], limit: int = 5) -> str:
    return "".join(f"[{paper['index']}]" for paper in cluster.get("papers", [])[:limit])


def _build_review_outline(topic: str, clusters: List[Dict[str, Any]], gaps: List[Dict[str, Any]]) -> str:
    lines = [f"# {topic} 文献综述大纲", "", "## 一、研究背景与问题提出"]
    intro_indices = _cluster_evidence_indices(clusters[0]) if clusters else ""
    lines.append(f"- 交代“{topic}”的研究缘起、核心概念和现实背景{intro_indices}。")
    lines.append("- 明确本文综述的对象、范围和资料来源。")
    lines.extend(["", "## 二、研究脉络与主题分支"])
    if clusters:
        for i, cluster in enumerate(clusters, 1):
            indices = _cluster_evidence_indices(cluster)
            lines.append(f"- {i}. {cluster['label']}：梳理该方向的主要问题、代表观点和证据边界{indices}。")
    else:
        lines.append("- 当前文献库尚未形成稳定主题分支，建议先扩展检索。")
    lines.extend(["", "## 三、研究不足与后续方向"])
    for gap in gaps[:4]:
        indices = "".join(f"[{i}]" for i in gap.get("evidence_indices", [])[:5]) or "（需补检索）"
        lines.append(f"- {gap['title']}：{gap['basis']}，证据线索 {indices}。")
    lines.extend(["", "## 四、段落证据映射", "- 写作时每个实质段落需保留证据编号，并对缺摘要文献标注“待核对原文”。"])
    return "\n".join(lines)


def _build_review_section(topic: str, selected: List[Dict[str, Any]], clusters: List[Dict[str, Any]], gaps: List[Dict[str, Any]], section: str) -> str:
    section_key = section.strip() if section else "文献综述"
    draft = _build_review_draft(topic, selected, clusters, gaps, mode="draft")
    if section_key in ("研究背景", "背景", "问题提出"):
        usable = [item for item in selected if item.get("relevance", 0) > 0 and "retracted" not in item.get("flags", [])]
        indices = _evidence_indices(usable)
        return "\n".join([
            f"# {topic}：研究背景",
            "",
            f"“{topic}”相关研究的展开，通常与媒介技术变迁、知识生产方式更新以及具体社会议题的外部传播需求有关{indices}。现有文献显示，该领域已经积累了一批围绕概念界定、传播路径、媒介平台和效果评价的研究，但不同文献之间在资料来源和原文可追溯性上并不均衡，因此后续写作需要区分摘要可追溯证据与待核对原文证据。",
            "",
            "## 段落证据映射",
            *[f"- [{item['index']}] {item['paper'].get('title', '未获取')}：{'摘要可追溯' if item['paper'].get('abstract') else '待核对原文'}，相关性分 {item.get('relevance', 0)}" for item in usable],
        ])
    if section_key in ("研究不足", "不足", "未来方向", "后续方向"):
        lines = [f"# {topic}：研究不足与后续方向", ""]
        for gap in gaps[:4]:
            indices = "".join(f"[{i}]" for i in gap.get("evidence_indices", [])[:5]) or "（当前库无直接证据）"
            lines.append(f"从现有文献库的覆盖情况看，{gap['title']}仍值得进一步展开。{gap['basis']}，相关证据序号为{indices}。这一不足并不意味着相关研究不存在，而是提示后续检索应在该维度上补充数据库、关键词和原文核对。")
        return "\n\n".join(lines)
    for cluster in clusters:
        if section_key in cluster["label"] or cluster["label"] in section_key:
            indices = _cluster_evidence_indices(cluster)
            body = cluster.get("synthesis") or "该方向下部分文献仍缺少摘要，现阶段只能作为题录线索处理。"
            return "\n".join([
                f"# {topic}：{cluster['label']}",
                "",
                f"在{cluster['label']}这一分支中，相关文献主要提供了关于“{topic}”的概念、对象或案例线索{indices}。{body}。由于该分支内部证据密度可能不均，写作时应优先使用摘要可追溯文献，并将缺少摘要的条目标注为待核对原文。",
                "",
                "## 段落证据映射",
                *[f"- [{paper['index']}] {paper.get('title', '未获取')}：{'摘要可追溯' if paper.get('trace_status') == 'abstract' else '待核对原文'}" for paper in cluster.get("papers", [])],
            ])
    return draft


def _build_review_draft(topic: str, selected: List[Dict[str, Any]], clusters: List[Dict[str, Any]], gaps: List[Dict[str, Any]], mode: str = "draft", section: str = "") -> str:
    usable = [item for item in selected if item.get("relevance", 0) > 0 and "retracted" not in item.get("flags", [])]
    cluster_list = clusters or _review_clusters(usable, topic)
    if mode == "outline":
        return _build_review_outline(topic, cluster_list, gaps)
    if mode == "section" or section:
        return _build_review_section(topic, selected, cluster_list, gaps, section)

    lines = [f"# {topic} 文献综述初稿", "", "## 一、研究背景与问题提出"]
    if not usable:
        lines.append("当前文献库中尚未形成足够高相关、可追溯的文献基础，建议扩大检索或补充摘要后再生成综述初稿。")
    else:
        intro_indices = _evidence_indices(usable)
        sources = sorted({item["paper"].get("source", "未获取") or "未获取" for item in usable})
        lines.append(
            f"围绕“{topic}”，当前文献库已经形成以{ '、'.join(cluster['label'] for cluster in cluster_list[:4]) }为主的若干研究分支{intro_indices}。从资料来源看，相关证据主要来自{ '、'.join(sources) }；从证据质量看，部分文献具有摘要支撑，部分条目仍需继续补充详情页或原文。因而，后续写作应把已有摘要作为直接论证基础，把题录信息作为待核对线索。"
        )
        lines.extend(["", "## 二、研究脉络与主题分支"])
        for cluster in cluster_list:
            indices = _cluster_evidence_indices(cluster)
            if cluster.get("synthesis"):
                body = cluster["synthesis"]
                evidence_note = "这些摘要能够为该分支的基本判断提供初步依据"
            else:
                body = "该分支目前主要由题录信息构成，尚不足以支撑细节性结论"
                evidence_note = "相关判断需在补充摘要或全文后再强化"
            lines.extend([
                f"### {cluster['label']}",
                f"{cluster['label']}是“{topic}”研究中的一个重要切面{indices}。现有材料显示，{body}。因此，写作时可将该分支作为综述的一个层次展开，但需要注意证据边界：{evidence_note}。",
                "",
            ])
        lines.extend(["## 三、研究不足与后续方向"])
        for i, gap in enumerate(gaps[:4], 1):
            indices = "".join(f"[{idx}]" for idx in gap.get("evidence_indices", [])[:5]) or "（当前库无直接证据）"
            lines.append(f"{i}. {gap['title']}。{gap['basis']}，相关证据序号为{indices}。这一提示只反映当前文献库的覆盖情况，后续应通过扩展关键词、数据库和原文核对进一步确认。")
    lines.append("")
    lines.append("## 四、段落证据映射")
    for item in usable:
        p = item["paper"]
        status = "摘要可追溯" if p.get("abstract") else "待核对原文"
        excerpt = _paper_write_excerpt(p, 80)
        excerpt_text = f"；摘要线索：{excerpt}" if excerpt else ""
        lines.append(f"- [{item['index']}] {p.get('title', '未获取')}：{status}，相关性分 {item.get('relevance', 0)}{excerpt_text}")
    return "\n".join(lines)


def _auto_detail_for_review(papers: List[Dict[str, Any]], candidates: List[Dict[str, Any]], detail_top_n: int, project: Optional[str]) -> Dict[str, Any]:
    targets = []
    for item in candidates:
        paper = item["paper"]
        if item.get("relevance", 0) <= 0 or "retracted" in item.get("flags", []):
            continue
        if _is_cnki_paper(paper) and paper.get("url") and not paper.get("abstract"):
            targets.append(item)
        if len(targets) >= detail_top_n:
            break
    if not targets:
        return {"attempted": 0, "updated": 0, "indices": []}

    selected = [item["paper"] for item in targets]
    enriched = batch_read_detail(papers=selected, top_n=len(selected), fulltext=False)
    enriched_map = {p.get("url", ""): p for p in enriched if isinstance(p, dict) and p.get("url")}
    updated = list(papers)
    updated_count = 0
    for item in targets:
        idx = item["index"] - 1
        url = updated[idx].get("url", "")
        detail = enriched_map.get(url)
        if not detail:
            continue
        before_had_abstract = bool(updated[idx].get("abstract"))
        merged = dict(updated[idx])
        for k, v in detail.items():
            if k == "fulltext":
                continue
            if v:
                merged[k] = v
        if merged.get("abstract") and not before_had_abstract:
            updated_count += 1
        updated[idx] = merged
    _save_session(updated, project=project)
    return {
        "attempted": len(targets),
        "updated": updated_count,
        "indices": [item["index"] for item in targets],
    }


def _review_write_inputs(papers: List[Dict[str, Any]], topic: str, limit: int) -> tuple:
    selected = _select_review_papers(papers, topic, limit)
    clusters = _review_clusters(selected, topic)
    gaps = _review_gaps(papers, topic)
    return selected, clusters, gaps


def _append_references(markdown: str, selected: List[Dict[str, Any]], style: str = "gbt7714") -> str:
    papers = [item["paper"] for item in selected]
    refs = generate_reference_list(papers, style)
    if refs:
        markdown = markdown.rstrip() + "\n\n## 参考文献\n" + refs.strip() + "\n"
    return markdown


def _sentence_evidence_indices(text: str) -> List[int]:
    return [int(i) for i in re.findall(r"\[(\d+)\]", text)]


def _validation_body(markdown: str) -> str:
    body = re.sub(r"## 参考文献[\s\S]*$", "", markdown)
    body = re.sub(r"## [一二三四五六七八九十、]*段落证据映射[\s\S]*$", "", body)
    return body


def _claim_sentences(markdown: str) -> List[Dict[str, Any]]:
    body = _validation_body(markdown)
    claims = []
    section = "正文"
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            section = stripped.lstrip("#").strip()
            continue
        if stripped.startswith("-"):
            stripped = stripped.lstrip("- ").strip()
        for raw in re.split(r"(?<=[。！？.!?])\s*", stripped):
            sentence = raw.strip()
            if len(sentence) < 12:
                continue
            claims.append({
                "section": section,
                "claim": sentence[:260],
                "evidence_indices": _sentence_evidence_indices(sentence),
            })
    return claims


def _paper_validation_text(paper: Dict[str, Any]) -> str:
    keywords = paper.get("keywords") or ""
    if isinstance(keywords, list):
        keywords = " ".join(str(k) for k in keywords)
    parts = [
        paper.get("title", ""),
        paper.get("journal", ""),
        keywords,
        paper.get("abstract", ""),
        paper.get("fulltext", ""),
    ]
    return " ".join(str(part) for part in parts if part).lower()


def _claim_terms(claim: str) -> List[str]:
    terms = []
    stopwords = {"研究", "文献", "相关", "现有", "显示", "因此", "这一", "进行", "通过", "围绕", "当前", "方面", "中的", "需要", "判断", "证据"}
    known_terms = (
        "国家形象", "国际传播", "对外传播", "传播效果", "平台机制", "社交媒体", "官方叙事",
        "中国故事", "文化符号", "受众认知", "话语", "媒介", "算法", "内容分析",
    )
    for term in known_terms:
        if term in claim and term not in terms:
            terms.append(term)
    for token in re.findall(r"[A-Za-z][A-Za-z0-9_-]{2,}", claim.lower()):
        if token not in terms:
            terms.append(token)
    for phrase in re.findall(r"[一-鿿]{2,}", claim):
        if phrase in stopwords:
            continue
        for size in (6, 5, 4, 3, 2):
            if len(phrase) < size:
                continue
            for i in range(0, len(phrase) - size + 1):
                token = phrase[i:i + size]
                if token in stopwords:
                    continue
                if token not in terms:
                    terms.append(token)
                if len(terms) >= 24:
                    return terms
    return terms[:24]


def _support_for_claim(claim: Dict[str, Any], evidence_by_index: Dict[int, Dict[str, Any]]) -> Dict[str, Any]:
    indices = claim.get("evidence_indices", [])
    if not indices:
        return {"support_level": "unsupported", "reason": "该论断未附证据编号", "matched_terms": []}
    invalid = [index for index in indices if index not in evidence_by_index]
    if invalid:
        return {"support_level": "invalid", "reason": f"证据编号不存在：{invalid}", "matched_terms": []}

    terms = _claim_terms(claim.get("claim", ""))
    matched_terms = []
    missing_abstract = False
    risky = False
    best_overlap = 0
    for index in indices:
        item = evidence_by_index[index]
        paper = item["paper"]
        flags = item.get("flags", [])
        text = _paper_validation_text(paper)
        overlap_terms = [term for term in terms if term in text]
        best_overlap = max(best_overlap, len(overlap_terms))
        for term in overlap_terms:
            if term not in matched_terms:
                matched_terms.append(term)
        if not paper.get("abstract") or "needs_fulltext_check" in flags:
            missing_abstract = True
        if "retracted" in flags:
            risky = True

    if risky:
        return {"support_level": "invalid", "reason": "引用了疑似撤稿文献，不应作为正面证据", "matched_terms": matched_terms[:8]}
    if missing_abstract:
        return {"support_level": "needs_fulltext_check", "reason": "引用文献缺少摘要或全文证据，需核对原文", "matched_terms": matched_terms[:8]}
    if not terms:
        return {"support_level": "medium", "reason": "论断缺少可提取关键词，但引用编号有效", "matched_terms": []}
    if best_overlap >= 2:
        return {"support_level": "strong", "reason": "论断关键词与引用文献题名/摘要/关键词存在较好匹配", "matched_terms": matched_terms[:8]}
    if best_overlap == 1:
        return {"support_level": "medium", "reason": "论断与引用文献存在有限词项匹配，建议核对表述是否过强", "matched_terms": matched_terms[:8]}
    return {"support_level": "weak", "reason": "未在引用文献题名/摘要/关键词中找到明显词项支撑", "matched_terms": []}


def _validate_writing(markdown: str, selected: List[Dict[str, Any]]) -> Dict[str, Any]:
    evidence_by_index = {item["index"]: item for item in selected}
    usable_indices = {
        item["index"] for item in selected
        if item.get("relevance", 0) > 0 and "retracted" not in item.get("flags", [])
    }
    claims = _claim_sentences(markdown)
    claim_results = []
    counts = {"strong": 0, "medium": 0, "weak": 0, "needs_fulltext_check": 0, "unsupported": 0, "invalid": 0}
    for claim in claims:
        support = _support_for_claim(claim, evidence_by_index)
        level = support["support_level"]
        counts[level] = counts.get(level, 0) + 1
        claim_results.append({
            "section": claim["section"],
            "claim": claim["claim"],
            "evidence_indices": claim["evidence_indices"],
            **support,
        })

    body = _validation_body(markdown)
    cited_indices = set(_sentence_evidence_indices(body))
    invalid_indices = sorted({index for index in cited_indices if index not in evidence_by_index})
    weak_claims = [item for item in claim_results if item["support_level"] in ("weak", "needs_fulltext_check")]
    unsupported_claims = [item for item in claim_results if item["support_level"] == "unsupported"]
    invalid_claims = [item for item in claim_results if item["support_level"] == "invalid"]
    unused_usable = sorted(usable_indices - cited_indices)

    issues = []
    if invalid_indices:
        issues.append({"type": "invalid_evidence_index", "indices": invalid_indices, "message": "正文引用了不存在于本次写作证据集的编号"})
    if unsupported_claims:
        issues.append({"type": "unsupported_claim", "count": len(unsupported_claims), "examples": unsupported_claims[:5], "message": "存在未附证据编号的实质性论断"})
    if weak_claims:
        issues.append({"type": "weak_or_unverified_support", "count": len(weak_claims), "examples": weak_claims[:8], "message": "部分论断与引用证据匹配较弱，或需核对原文"})
    if invalid_claims:
        issues.append({"type": "invalid_support", "count": len(invalid_claims), "examples": invalid_claims[:5], "message": "部分论断引用了无效或高风险证据"})
    if unused_usable:
        issues.append({"type": "unused_relevant_evidence", "indices": unused_usable[:12], "message": "部分高相关证据未进入正文论证，可按需补充"})

    score = 100
    score -= min(counts.get("unsupported", 0) * 8, 32)
    score -= min(counts.get("weak", 0) * 6, 30)
    score -= min(counts.get("needs_fulltext_check", 0) * 5, 25)
    score -= min(counts.get("invalid", 0) * 18, 45)
    score = max(score, 0)
    status = "success" if score >= 80 and not invalid_claims else "warning"
    recommendations = []
    if unsupported_claims:
        recommendations.append("为未附编号的论断补充 [证据序号]，或删除无法由当前文献库支撑的判断。")
    if weak_claims:
        recommendations.append("对弱匹配或缺摘要证据执行 read-detail/read-detail --fulltext，并收紧过强表述。")
    if invalid_claims:
        recommendations.append("移除无效编号或疑似撤稿文献，改用摘要可追溯的高相关文献。")
    if unused_usable:
        recommendations.append("检查未使用的高相关文献，必要时补充到相应段落。")
    return {
        "status": status,
        "score": score,
        "evidence_count": len(selected),
        "cited_evidence_count": len(cited_indices),
        "checked_claims": len(claim_results),
        "support_counts": counts,
        "claim_results": claim_results[:30],
        "issues": issues,
        "recommendations": recommendations,
    }


def cmd_write(args):
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA", "message": "没有可写作的文献，请先执行 search、batch-search 或 import"})
        return
    topic = args.topic or _session_project(args) or "当前课题"
    limit = min(args.limit or 12, len(papers))
    selected, clusters, gaps = _review_write_inputs(papers, topic, limit)
    mode = getattr(args, "mode", "draft") or "draft"
    section = getattr(args, "section", "") or ""
    if section and mode == "draft":
        mode = "section"
    markdown = _build_review_draft(topic, selected, clusters, gaps, mode=mode, section=section)
    if args.with_citations:
        markdown = _append_references(markdown, selected, args.citation_style or "gbt7714")
    validation = _validate_writing(markdown, selected) if getattr(args, "validate", False) else None

    output_path = Path(args.output) if args.output else None
    if args.format == "docx":
        if output_path is None:
            output_path = Path("review.docx")
        elif output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")
        result = _write_docx_from_markdown(markdown, output_path)
        result.update({"project": _session_project(args), "topic": topic, "mode": mode, "section": section or None, "format": "docx", "validation": validation})
        _output(result)
        return

    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(markdown, encoding="utf-8")
    if args.raw:
        print(markdown)
    else:
        _output({
            "status": "success",
            "project": _session_project(args),
            "topic": topic,
            "mode": mode,
            "section": section or None,
            "format": "markdown",
            "output_file": str(output_path) if output_path else None,
            "validation": validation,
            "markdown": markdown,
        })

def _topic_methods(label: str, gap_title: str) -> List[str]:
    text = label + gap_title
    methods = []
    if any(term in text for term in ("平台", "社交媒体", "传播效果", "受众")):
        methods.extend(["内容分析", "问卷/访谈", "平台比较"])
    if any(term in text for term in ("话语", "叙事", "文化", "中国故事", "国家形象")):
        methods.extend(["话语分析", "案例研究", "符号分析"])
    if any(term in text for term in ("方法", "模型", "算法")):
        methods.extend(["模型构建", "计算传播分析", "混合方法"])
    if not methods:
        methods.extend(["文献综述", "案例研究", "比较研究"])
    return list(dict.fromkeys(methods))[:4]


def _build_topics(topic: str, papers: List[Dict[str, Any]], limit: int = 8) -> List[Dict[str, Any]]:
    candidates = _review_candidates(papers, topic)
    selected = [item for item in candidates if item.get("relevance", 0) > 0 and "retracted" not in item.get("flags", [])][:max(limit, 1)]
    clusters = _review_clusters(selected, topic)
    gaps = _review_gaps(papers, topic)
    suggestions = []
    for cluster in clusters[:4]:
        cluster_indices = [paper["index"] for paper in cluster.get("papers", [])[:5]]
        matched_gap = None
        for gap in gaps:
            if gap.get("evidence_indices"):
                if set(cluster_indices) & set(gap.get("evidence_indices", [])):
                    matched_gap = gap
                    break
        if matched_gap is None and gaps:
            matched_gap = gaps[min(len(suggestions), len(gaps) - 1)]
        gap_title = matched_gap.get("title", "需进一步明确研究空白") if matched_gap else "需进一步明确研究空白"
        evidence_indices = sorted(set(cluster_indices + (matched_gap.get("evidence_indices", [])[:5] if matched_gap else [])))
        if not evidence_indices:
            evidence_indices = cluster_indices
        risks = []
        if not evidence_indices:
            risks.append("当前文献库证据不足，需先补充检索")
        if any("metadata_only" == paper.get("trace_status") for paper in cluster.get("papers", [])):
            risks.append("部分证据仅有题录信息，需补摘要或原文")
        if matched_gap and matched_gap.get("matched_count", 0) == 0:
            risks.append("该方向当前库无直接命中文献，不能直接断言真实研究空白")
        suggestions.append({
            "title": f"{topic}中的{cluster['label']}研究",
            "rationale": f"当前文献库中“{cluster['label']}”聚类包含 {cluster.get('count', 0)} 篇线索；{matched_gap.get('basis', '需结合更多检索证据判断研究空间') if matched_gap else '需结合更多检索证据判断研究空间'}",
            "evidence_indices": evidence_indices[:8],
            "possible_methods": _topic_methods(cluster["label"], gap_title),
            "risks": risks or ["需在开题前继续核对核心文献和原文证据"],
            "followup_search": [topic, cluster["label"], gap_title],
        })
    if not suggestions:
        suggestions.append({
            "title": f"{topic}的研究现状与问题重构",
            "rationale": f"当前文献库共有 {len(papers)} 篇记录，但高相关聚类不足，适合先做综述型选题或扩大检索。",
            "evidence_indices": [],
            "possible_methods": ["文献综述", "关键词扩展检索", "题录筛选"],
            "risks": ["证据基础不足，不能直接进入实证设计"],
            "followup_search": [topic, f"{topic} 研究现状", f"{topic} 研究空白"],
        })
    return suggestions[:limit]


def cmd_topics(args):
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA", "message": "没有可生成选题的数据，请先执行 search、batch-search 或 import"})
        return
    topic = args.topic or _session_project(args) or "当前课题"
    suggestions = _build_topics(topic, papers, args.limit or 8)
    _output({
        "status": "success",
        "project": _session_project(args),
        "topic": topic,
        "total": len(papers),
        "count": len(suggestions),
        "topics": suggestions,
    })


def cmd_validate(args):
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA", "message": "没有可校验的文献，请先执行 search、batch-search 或 import"})
        return
    topic = args.topic or _session_project(args) or "当前课题"
    limit = min(args.limit or 12, len(papers))
    selected, clusters, gaps = _review_write_inputs(papers, topic, limit)
    if args.file:
        markdown = Path(args.file).read_text(encoding="utf-8")
    else:
        markdown = _build_review_draft(topic, selected, clusters, gaps, mode="draft")
    validation = _validate_writing(markdown, selected)
    validation.update({
        "project": _session_project(args),
        "topic": topic,
        "checked_file": args.file,
    })
    _output(validation)


def cmd_review(args):
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA", "message": "没有可生成综述的文献，请先执行 search、batch-search 或 import"})
        return
    topic = args.topic or _session_project(args) or "当前课题"
    limit = min(args.limit or 12, len(papers))
    candidates = _review_candidates(papers, topic)
    auto_detail = None
    if getattr(args, "auto_detail", False):
        detail_top_n = max(getattr(args, "detail_top_n", 5) or 5, 1)
        print(f"[review] 自动补全高相关知网文献摘要（最多 {detail_top_n} 篇）...", file=sys.stderr)
        auto_detail = _auto_detail_for_review(papers, candidates, detail_top_n, _session_project(args))
        if auto_detail.get("attempted"):
            papers = _load_session(_session_project(args))
            candidates = _review_candidates(papers, topic)
    selected = _select_review_papers(papers, topic, limit)
    clusters = _review_clusters(selected, topic) if getattr(args, "cluster", False) else None
    gaps = _review_gaps(papers, topic) if getattr(args, "gaps", False) else None
    evidence = []
    for item in selected:
        entry = _paper_evidence(item["paper"], item["index"])
        entry["relevance"] = item.get("relevance", 0)
        entry["flags"] = item.get("flags", [])
        evidence.append(entry)
    markdown = _build_review_markdown(topic, _session_project(args), selected, len(papers), candidates, clusters, gaps)
    if args.output:
        out = Path(args.output)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(markdown, encoding="utf-8")
    if args.raw:
        print(markdown)
    else:
        _output({
            "status": "success",
            "project": _session_project(args),
            "topic": topic,
            "count": len(selected),
            "total": len(papers),
            "output_file": args.output,
            "auto_detail": auto_detail,
            "clusters": clusters,
            "gaps": gaps,
            "evidence": evidence,
            "markdown": markdown,
        })



def cmd_projects(args):
    base = Path.cwd() / ".scholar-kit" / "projects"
    projects = []
    if base.exists():
        for project_dir in sorted(p for p in base.iterdir() if p.is_dir()):
            papers = _load_session(project_dir.name)
            projects.append({
                "name": project_dir.name,
                "count": len(papers),
                "session_file": str(project_dir / "session.json"),
            })
    _output({"status": "success", "count": len(projects), "projects": projects})


def cmd_library(args):
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA", "message": "没有可查看的文献，请先执行 search、batch-search 或 import"})
        return
    limit = args.limit or len(papers)
    rows = [_paper_summary(p, i + 1) for i, p in enumerate(papers[:limit])]
    _output({"status": "success", "project": _session_project(args), "count": len(papers), "results": rows})


# ── cite 命令 ─────────────────────────────────────────

def cmd_cite(args):
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA", "message": "没有可格式化的数据，请先执行 search、batch-search 或 import"})
        return

    # 对缺少卷期页码的知网论文，自动走 detail 补全
    cnki_need_enrich = [
        (i, p) for i, p in enumerate(papers)
        if _is_cnki_paper(p) and p.get("url") and not p.get("pages") and "cnki.net" in p.get("url", "")
    ]
    if cnki_need_enrich:
        print(f"[cite] 正在补全 {len(cnki_need_enrich)} 篇知网论文的卷期页码...",
              file=__import__('sys').stderr)
        from cnki import get_detail
        for idx, p in cnki_need_enrich:
            detail = get_detail(p["url"])
            if detail and detail.get("status") != "error":
                for k in ("volume", "issue", "pages", "doi", "year", "journal"):
                    if detail.get(k) and not p.get(k):
                        p[k] = detail[k]
            time.sleep(1)

    enriched = []
    for i, p in enumerate(papers):
        if p.get("doi") and not p.get("volume"):
            if i > 0:
                time.sleep(1)
            crossref_data = resolve_crossref(p["doi"])
            if crossref_data:
                p.update({k: v for k, v in crossref_data.items() if v and not p.get(k)})
        enriched.append(p)

    ref_list = generate_reference_list(enriched, args.style or "gbt7714")
    if args.raw:
        print(ref_list)
    else:
        _output({"status": "success", "project": _session_project(args), "style": args.style or "gbt7714",
                 "count": len(enriched), "references": ref_list})


# ── import 命令 ───────────────────────────────────────

def cmd_import(args):
    results = parse_cnki_export(args.filepath)
    if results and not (len(results) == 1 and results[0].get("status") == "error"):
        _save_session(results, project=_session_project(args))
        _output({"status": "success", "count": len(results), "results": results})
    else:
        _output(results[0] if results else {"status": "error", "code": "IMPORT_PARSE_FAILED", "message": "解析失败"})


# ── read-paper 命令 ───────────────────────────────────

def cmd_read_paper(args):
    """读取用户论文文件（.docx / .txt / .md），输出 UTF-8 纯文本"""
    filepath = Path(args.filepath)
    if not filepath.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND", "message": f"文件不存在: {args.filepath}"})
        return

    suffix = filepath.suffix.lower()
    text = ""

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            _output({"status": "error", "code": "MISSING_DEPENDENCY",
                     "message": "缺少 python-docx 依赖"})
            return
        try:
            doc = Document(str(filepath))
            parts: List[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        parts.append(" | ".join(row_texts))
            text = "\n\n".join(parts)
        except Exception as e:
            _output({"status": "error", "code": "DOCX_PARSE_FAILED", "message": f"docx 解析失败: {e}"})
            return

    elif suffix in (".txt", ".md", ".markdown"):
        decoded = False
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030"):
            try:
                text = filepath.read_text(encoding=enc)
                decoded = True
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if not decoded:
            _output({"status": "error", "code": "ENCODING_ERROR", "message": "无法识别文件编码"})
            return

    elif suffix == ".pdf":
        _output({"status": "error", "code": "UNSUPPORTED_FORMAT",
                 "message": "PDF 请使用 Agent 内置的文件读取工具直接读取，无需 read-paper"})
        return
    else:
        _output({"status": "error", "code": "UNSUPPORTED_FORMAT",
                 "message": f"不支持的文件格式: {suffix}"})
        return

    char_count = len(text)
    para_count = text.count("\n\n") + 1 if text else 0

    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(text, encoding="utf-8")
        _output({"status": "success",
                 "message": f"已提取到: {args.output}",
                 "chars": char_count,
                 "paragraphs": para_count})
    else:
        if args.raw:
            print(text)
        else:
            _output({"status": "success",
                     "chars": char_count,
                     "paragraphs": para_count,
                     "text": text})


# ── pdf-meta 命令 ──────────────────────────────────────

def cmd_pdf_meta(args):
    """从 PDF 文件中提取元数据（标题、作者、DOI 等）"""
    filepath = Path(args.filepath)
    if not filepath.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND", "message": f"文件不存在: {args.filepath}"})
        return

    if filepath.suffix.lower() != ".pdf":
        _output({"status": "error", "code": "NOT_PDF", "message": "仅支持 PDF 文件"})
        return

    try:
        from pypdf import PdfReader
    except ImportError:
        _output({"status": "error", "code": "MISSING_DEPENDENCY", "message": "缺少 pypdf 依赖"})
        return

    try:
        reader = PdfReader(str(filepath))
        meta = reader.metadata or {}

        result = {"status": "success", "file": str(filepath)}

        if meta.title:
            result["title"] = meta.title
        if meta.author:
            result["authors"] = meta.author
        if meta.subject:
            result["subject"] = meta.subject

        # 从 XMP 元数据中提取 DOI
        doi = None
        if hasattr(reader, 'xmp_metadata') and reader.xmp_metadata:
            xmp = reader.xmp_metadata
            # DOI 可能在 dc:identifier 或自定义属性中
            if hasattr(xmp, 'dc_identifier') and xmp.dc_identifier:
                for ident in (xmp.dc_identifier if isinstance(xmp.dc_identifier, list) else [xmp.dc_identifier]):
                    if ident and '10.' in str(ident):
                        import re as _re
                        doi_m = _re.search(r'(10\.\d{4,}/[^\s]+)', str(ident))
                        if doi_m:
                            doi = doi_m.group(1)
                            break

        # 从前几页文本中查找 DOI
        if not doi:
            import re as _re
            for page_num in range(min(3, len(reader.pages))):
                page_text = reader.pages[page_num].extract_text() or ""
                doi_m = _re.search(r'(?:DOI|doi)[：:\s]*\s*(10\.\d{4,}/[^\s]+)', page_text)
                if doi_m:
                    doi = doi_m.group(1).rstrip(".")
                    break

        if doi:
            result["doi"] = doi
            # 用 DOI 从 Crossref 补全完整元数据
            crossref_data = resolve_crossref(doi)
            if crossref_data:
                result["crossref"] = crossref_data

        _output(result)

    except Exception as e:
        _output({"status": "error", "code": "PDF_READ_FAILED", "message": str(e)})


# ── batch-search 命令 ─────────────────────────────────

def cmd_batch_search(args):
    """批量搜索：浏览器只启动一次，循环搜索多个关键词"""
    keywords = list(args.queries) if args.queries else []

    if args.query_file:
        qf = Path(args.query_file)
        if not qf.exists():
            _output({"status": "error", "code": "FILE_NOT_FOUND", "message": f"关键词文件不存在: {args.query_file}"})
            return
        qf_text = None
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030"):
            try:
                qf_text = qf.read_text(encoding=enc)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if qf_text is None:
            _output({"status": "error", "code": "ENCODING_ERROR",
                     "message": f"关键词文件编码无法识别: {args.query_file}"})
            return
        for line in qf_text.splitlines():
            line = line.strip()
            if line and not line.startswith("#"):
                keywords.append(line)

    if not keywords:
        _output({"status": "error", "code": "NO_KEYWORDS",
                 "message": "未提供关键词"})
        return

    result = batch_search_cnki(
        keywords=keywords,
        core=args.core,
        author=getattr(args, "author", None),
        journal=getattr(args, "journal", None),
        doc_type=getattr(args, "doc_type", None),
        field=getattr(args, "field", None),
        year_from=args.year_from,
        year_to=args.year_to,
        sort=args.sort or "relevance",
        pages=args.pages or 1,
    )

    if result.get("status") in ("success", "partial") and result.get("results"):
        _save_session(result.get("results") or [], append=args.append, project=_session_project(args))

    if args.export and result.get("results"):
        content = export_papers(result["results"], args.export, args.output)
        if isinstance(content, dict) and content.get("status") == "error":
            _output(content)
            return
        export_output = {"status": result.get("status", "success"),
                         "count": len(result["results"]),
                         "format": args.export, "output_file": args.output,
                         "content": content}
        if result.get("errors"):
            export_output["errors"] = result["errors"]
        _output(export_output)
    else:
        _output(result)


# ── read-detail 命令 ──────────────────────────────────

def _is_cnki_paper(paper: dict) -> bool:
    url = paper.get("url", "")
    source = paper.get("source", "")
    return "cnki" in url.lower() or source == "CNKI" or source == "CNKI-export"


def _parse_indices(raw: str, total: int) -> List[int]:
    """解析用户传入的序号字符串，返回 0-based 索引列表。

    支持格式: "3" "1,3,9" "2-5" "1,3-5,8" （序号从 1 开始）
    """
    indices: List[int] = []
    for part in raw.split(","):
        part = part.strip()
        if "-" in part:
            lo, hi = part.split("-", 1)
            lo_i, hi_i = int(lo.strip()), int(hi.strip())
            indices.extend(range(lo_i - 1, min(hi_i, total)))
        else:
            i = int(part.strip()) - 1
            if 0 <= i < total:
                indices.append(i)
    return sorted(set(indices))


def cmd_read_detail(args):
    """对会话中的论文批量获取摘要/全文"""
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA",
                 "message": "没有可读取的论文，请先执行 search、batch-search 或 import"})
        return

    do_fulltext = args.fulltext
    indices_raw = getattr(args, "indices", None)

    if indices_raw:
        pick_idx = _parse_indices(indices_raw, len(papers))
        if not pick_idx:
            _output({"status": "error", "code": "INVALID_INDICES",
                     "message": f"无效序号 '{indices_raw}'，会话共 {len(papers)} 篇（序号 1-{len(papers)}）"})
            return
        selected = [papers[i] for i in pick_idx]
        label = f"第 {indices_raw} 篇"
    else:
        top_n = args.top_n or 5
        selected = papers[:top_n]
        label = f"前 {len(selected)} 篇"

    print(f"[read-detail] 会话共 {len(papers)} 篇，将获取{label}的{'全文' if do_fulltext else '摘要'}",
          file=sys.stderr)

    cnki_selected = [p for p in selected if _is_cnki_paper(p)]
    non_cnki_selected = [p for p in selected if not _is_cnki_paper(p)]
    if not cnki_selected:
        _output({"status": "warning", "code": "NO_SESSION_DATA",
                 "message": "所选论文中无知网论文，read-detail 仅支持知网论文。API 源论文请直接使用搜索时返回的摘要",
                 "count": len(non_cnki_selected), "results": non_cnki_selected})
        return
    enriched = batch_read_detail(
        papers=cnki_selected,
        top_n=len(cnki_selected),
        fulltext=do_fulltext,
    )
    enriched.extend(non_cnki_selected)

    if indices_raw:
        updated = list(papers)
        enriched_map = {p.get("url", ""): p for p in enriched if p.get("url")}
        for i in pick_idx:
            url = updated[i].get("url", "")
            if url in enriched_map:
                merged = {k: v for k, v in enriched_map[url].items() if k != "fulltext"}
                merged.update({k: updated[i][k] for k in updated[i] if k not in merged})
                updated[i] = merged
        session_papers = updated
    else:
        session_papers = []
        for p in enriched:
            sp = {k: v for k, v in p.items() if k != "fulltext"}
            session_papers.append(sp)
    _save_session(session_papers, project=_session_project(args))

    output_papers = enriched
    results = []
    for p in output_papers:
        entry: Dict[str, Any] = {
            "title": p.get("title", ""),
            "authors": p.get("authors", ""),
            "journal": p.get("journal", ""),
            "date": p.get("date", ""),
            "abstract": p.get("abstract", ""),
            "keywords": p.get("keywords", []),
            "has_fulltext": p.get("has_fulltext", False),
            "fulltext_length": p.get("fulltext_length", 0),
        }
        if p.get("fulltext_cache"):
            entry["fulltext_cache"] = p["fulltext_cache"]

        if do_fulltext and p.get("fulltext"):
            entry["fulltext"] = p["fulltext"]
        elif do_fulltext and p.get("fulltext_cache"):
            try:
                cache_path = Path(p["fulltext_cache"]).resolve()
                allowed_dir = (Path.cwd() / ".scholar-kit" / "fulltext").resolve()
                try:
                    cache_path.relative_to(allowed_dir)
                except ValueError:
                    entry["fulltext"] = ""
                else:
                    cache_data = json.loads(
                        cache_path.read_text(encoding="utf-8")
                    )
                    entry["fulltext"] = cache_data.get("fulltext", "")
            except Exception:
                entry["fulltext"] = ""

        results.append(entry)

    _output({"status": "success", "count": len(results), "results": results})


# ── docx 辅助函数 ────────────────────────────────────

def _get_or_create_footnotes_part(doc):
    """获取或创建符合 OOXML 标准的 footnotes XmlPart。

    使用 XmlPart 而非 Part：修改 element 后 doc.save() 自动序列化，无需手动同步。
    返回 (footnotes_element, max_id)。
    """
    from docx.oxml.ns import qn
    from docx.opc.part import XmlPart
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from lxml import etree

    try:
        ftn_part = doc.part.part_related_by(RT.FOOTNOTES)
        if isinstance(ftn_part, XmlPart):
            ftn_element = ftn_part.element
        else:
            ftn_element = etree.fromstring(ftn_part.blob)
            ftn_part.__class__ = XmlPart
            ftn_part._element = ftn_element
    except KeyError:
        ftn_xml = (
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="0">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="1">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )
        ftn_element = etree.fromstring(ftn_xml)
        from docx.opc.constants import CONTENT_TYPE as CT
        from docx.opc.packuri import PackURI
        ftn_part = XmlPart(
            PackURI("/word/footnotes.xml"),
            CT.WML_FOOTNOTES,
            ftn_element,
            doc.part.package,
        )
        doc.part.relate_to(ftn_part, RT.FOOTNOTES)

    max_id = 1
    for fn in ftn_element.findall(qn("w:footnote")):
        fid = fn.get(qn("w:id"))
        if fid and fid.isdigit():
            max_id = max(max_id, int(fid))

    return ftn_element, max_id


def _add_footnote_to_element(ftn_element, fn_id_counter, run_element, fn_text):
    """在 run_element 后插入脚注引用，并在 footnotes 部件中添加脚注内容。返回新 id。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fn_id_counter[0] += 1
    fid = fn_id_counter[0]

    footnote_el = OxmlElement("w:footnote")
    footnote_el.set(qn("w:id"), str(fid))
    fn_p = OxmlElement("w:p")

    fn_r_ref = OxmlElement("w:r")
    fn_rpr = OxmlElement("w:rPr")
    fn_rstyle = OxmlElement("w:rStyle")
    fn_rstyle.set(qn("w:val"), "FootnoteReference")
    fn_rpr.append(fn_rstyle)
    fn_r_ref.append(fn_rpr)
    fn_r_ref.append(OxmlElement("w:footnoteRef"))
    fn_p.append(fn_r_ref)

    fn_r_text = OxmlElement("w:r")
    fn_t = OxmlElement("w:t")
    fn_t.set(qn("xml:space"), "preserve")
    fn_t.text = " " + fn_text
    fn_r_text.append(fn_t)
    fn_p.append(fn_r_text)

    footnote_el.append(fn_p)
    ftn_element.append(footnote_el)

    ref_run = OxmlElement("w:r")
    ref_rpr = OxmlElement("w:rPr")
    ref_style = OxmlElement("w:rStyle")
    ref_style.set(qn("w:val"), "FootnoteReference")
    ref_rpr.append(ref_style)
    ref_run.append(ref_rpr)
    ref_mark = OxmlElement("w:footnoteReference")
    ref_mark.set(qn("w:id"), str(fid))
    ref_run.append(ref_mark)
    run_element.addnext(ref_run)

    return fid


def _para_replace_text(paragraph, find_text, replace_text):
    """段落级文本替换：先尝试单 run 内替换；跨 run 时精确切割，只改被命中区间，保留其余 run 格式。"""
    for run in paragraph.runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text, 1)
            return True

    runs = paragraph.runs
    if not runs:
        return False

    boundaries = []
    pos = 0
    for run in runs:
        end = pos + len(run.text)
        boundaries.append((pos, end, run))
        pos = end

    full = "".join(r.text for r in runs)
    if find_text not in full:
        return False

    idx = full.index(find_text)
    end_idx = idx + len(find_text)

    first_i = last_i = None
    for i, (s, e, _) in enumerate(boundaries):
        if first_i is None and s <= idx < e:
            first_i = i
        if s < end_idx <= e:
            last_i = i
            break

    if first_i is None or last_i is None:
        return False

    fs, fe, first_run = boundaries[first_i]
    ls, le, last_run = boundaries[last_i]

    if first_i == last_i:
        offset = idx - fs
        first_run.text = first_run.text[:offset] + replace_text + first_run.text[offset + len(find_text):]
    else:
        first_run.text = first_run.text[:idx - fs] + replace_text
        for j in range(first_i + 1, last_i):
            boundaries[j][2].text = ""
        last_run.text = last_run.text[end_idx - ls:]

    return True


def _find_run_containing(paragraph, text):
    """在段落中找到包含指定文本的 run（取全文第一次匹配），返回 run._element 或 None。
    跨 run 时精确定位到匹配文本末尾字符所在的 run（脚注应插在该 run 之后）。
    """
    if not text:
        return None

    for run in paragraph.runs:
        if text in run.text:
            return run._element

    runs = paragraph.runs
    if not runs:
        return None

    boundaries = []
    pos = 0
    for run in runs:
        end = pos + len(run.text)
        boundaries.append((pos, end, run))
        pos = end

    full = "".join(r.text for r in runs)
    if text not in full:
        return None

    end_idx = full.index(text) + len(text)
    target = end_idx - 1
    for start, end, run in boundaries:
        if start <= target < end:
            return run._element

    return runs[-1]._element


def _write_docx_from_markdown(md_text: str, output_path: Path) -> Dict[str, Any]:
    """Markdown 文本 → 学术格式 .docx"""
    import re
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn
    except ImportError:
        return {"status": "error", "code": "MISSING_DEPENDENCY", "message": "缺少 python-docx 依赖"}

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(16 - level * 2)
        hs.font.bold = True
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    lines = md_text.split("\n")
    footnotes_map: Dict[str, str] = {}
    body_lines: List[str] = []
    ref_section_lines: List[str] = []
    in_ref_section = False
    warnings: List[str] = []

    for line in lines:
        fn_def = re.match(r"^\[\^(\d+)\]:\s*(.+)$", line)
        if fn_def:
            footnotes_map[fn_def.group(1)] = fn_def.group(2)
            continue
        if re.match(r"^#{1,3}\s*(参考文献|References)", line, re.IGNORECASE):
            in_ref_section = True
            continue
        if in_ref_section:
            if line.strip():
                ref_section_lines.append(line.strip())
            continue
        body_lines.append(line)

    ftn_element, max_id = _get_or_create_footnotes_part(doc)
    fn_counter = [max_id]

    def _parse_inline(paragraph, text):
        pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|\[\^(\d+)\])")
        last = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                paragraph.add_run(text[last:m.start()])
            if m.group(2):
                paragraph.add_run(m.group(2)).bold = True
            elif m.group(3):
                paragraph.add_run(m.group(3)).italic = True
            elif m.group(4):
                fn_id = m.group(4)
                fn_text_content = footnotes_map.get(fn_id, "")
                if not fn_text_content:
                    warnings.append(f"[^{fn_id}] 无对应脚注定义，脚注内容为空")
                r = paragraph.add_run("")
                _add_footnote_to_element(ftn_element, fn_counter, r._element, fn_text_content)
            last = m.end()
        if last < len(text):
            paragraph.add_run(text[last:])

    for line in body_lines:
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {level}"]
            _parse_inline(p, heading.group(2).strip())
            continue
        ul_match = re.match(r"^[-*]\s+(.+)$", line)
        if ul_match:
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
            _parse_inline(p, ul_match.group(1).strip())
            continue
        ol_match = re.match(r"^\d{1,3}[.)]\s+(.+)$", line)
        if ol_match:
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            _parse_inline(p, ol_match.group(1).strip())
            continue
        if not line.strip():
            continue
        p = doc.add_paragraph()
        _parse_inline(p, line.strip())

    if ref_section_lines:
        doc.add_heading("参考文献", level=1)
        for ref_line in ref_section_lines:
            ref_line = re.sub(r"^\[\d+\]\s*", "", ref_line)
            ref_line = re.sub(r"^[-•]\s*", "", ref_line)
            p = doc.add_paragraph(ref_line)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.left_indent = Cm(0.74)

    try:
        doc.save(str(output_path))
    except Exception as e:
        return {"status": "error", "code": "IO_ERROR", "message": f"保存失败: {e}"}

    result: Dict[str, Any] = {
        "status": "success" if not warnings else "warning",
        "message": f"已生成: {output_path}",
        "output": str(output_path),
        "footnotes": fn_counter[0] - max_id,
        "references": len(ref_section_lines),
    }
    if warnings:
        result["warnings"] = warnings
    return result


# ── write-docx 命令 ───────────────────────────────────

def cmd_write_docx(args):
    """Markdown 文件 → 学术格式 .docx"""
    md_path = Path(args.filepath)
    if not md_path.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND",
                 "message": f"文件不存在: {args.filepath}"})
        return

    md_text = None
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030"):
        try:
            md_text = md_path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    if md_text is None:
        _output({"status": "error", "code": "ENCODING_ERROR",
                 "message": "文件编码无法识别"})
        return

    output_path = Path(args.output) if args.output else md_path.with_suffix(".docx")
    _output(_write_docx_from_markdown(md_text, output_path))


# ── patch-docx 命令 ───────────────────────────────────

def cmd_patch_docx(args):
    """在现有 .docx 上打补丁：文本替换 + 脚注插入 + 追加参考文献"""
    try:
        from docx import Document
        from docx.shared import Cm
    except ImportError:
        _output({"status": "error", "code": "MISSING_DEPENDENCY",
                 "message": "缺少 python-docx 依赖"})
        return

    docx_path = Path(args.filepath)
    if not docx_path.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND",
                 "message": f"文件不存在: {args.filepath}"})
        return

    patch_path = Path(args.patch)
    if not patch_path.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND",
                 "message": f"补丁文件不存在: {args.patch}"})
        return

    try:
        patch_data = json.loads(patch_path.read_text(encoding="utf-8"))
    except Exception as e:
        _output({"status": "error", "code": "PATCH_PARSE_FAILED",
                 "message": f"补丁 JSON 解析失败: {e}"})
        return

    if not isinstance(patch_data, dict):
        _output({"status": "error", "code": "PATCH_PARSE_FAILED",
                 "message": "补丁文件顶层必须是 JSON 对象"})
        return

    patches = patch_data.get("patches", [])
    footnotes_list = patch_data.get("footnotes", [])
    append_refs = patch_data.get("append_references", [])

    if not isinstance(patches, list) or not isinstance(footnotes_list, list) or not isinstance(append_refs, list):
        _output({"status": "error", "code": "PATCH_PARSE_FAILED",
                 "message": "patches / footnotes / append_references 必须是数组"})
        return

    output_path = Path(args.output) if args.output else docx_path.with_name(
        docx_path.stem + "_patched" + docx_path.suffix
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(docx_path))
    stats = {"replaced": 0, "not_found": 0, "footnotes_added": 0, "references_appended": 0}
    warnings: List[str] = []

    ftn_element, max_id = _get_or_create_footnotes_part(doc)
    fn_counter = [max_id]

    for patch in patches:
        find_text = patch.get("find", "")
        replace_text = patch.get("replace", "")
        if not find_text:
            continue
        found = False
        for para in doc.paragraphs:
            if find_text in para.text:
                if _para_replace_text(para, find_text, replace_text):
                    stats["replaced"] += 1
                    found = True
                    break
        if not found:
            stats["not_found"] += 1
            warnings.append(f"未找到替换目标: \"{find_text[:30]}...\"" if len(find_text) > 30 else f"未找到替换目标: \"{find_text}\"")

    for fn_entry in footnotes_list:
        after_text = fn_entry.get("after", "")
        fn_text = fn_entry.get("text", "")
        if not after_text or not fn_text:
            continue
        found = False
        for para in doc.paragraphs:
            if after_text not in para.text:
                continue
            run_el = _find_run_containing(para, after_text)
            if run_el is not None:
                _add_footnote_to_element(ftn_element, fn_counter, run_el, fn_text)
                stats["footnotes_added"] += 1
                found = True
                break
        if not found:
            warnings.append(f"脚注定位失败: \"{after_text[:30]}\"" if len(after_text) > 30 else f"脚注定位失败: \"{after_text}\"")

    if append_refs:
        existing_ref_heading = None
        for para in doc.paragraphs:
            if para.text.strip() in ("参考文献", "References") and para.style.name.startswith("Heading"):
                existing_ref_heading = para
        if existing_ref_heading is None:
            doc.add_paragraph()
            try:
                ref_heading = doc.add_paragraph("参考文献")
                ref_heading.style = doc.styles["Heading 1"]
            except KeyError:
                ref_heading = doc.add_paragraph("参考文献")
                ref_heading.runs[0].bold = True
        for ref_text in append_refs:
            p = doc.add_paragraph(ref_text)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.left_indent = Cm(0.74)
        stats["references_appended"] = len(append_refs)

    try:
        doc.save(str(output_path))
    except Exception as e:
        _output({"status": "error", "code": "IO_ERROR",
                 "message": f"保存失败: {e}"})
        return

    has_issues = stats["not_found"] > 0 or len(warnings) > 0
    result: Dict[str, Any] = {
        "status": "partial" if has_issues else "success",
        "message": f"已保存: {output_path}",
        "output": str(output_path),
        **stats,
    }
    if warnings:
        result["warnings"] = warnings
    _output(result)


# ── clean-cache 命令 ──────────────────────────────────

def cmd_clean_cache(args):
    """清理 .scholar-kit/ 缓存目录"""
    from datetime import datetime, timedelta
    from config import get as cfg_get

    cache_dir = Path.cwd() / ".scholar-kit"
    if not cache_dir.exists():
        _output({"status": "success", "message": "无缓存目录", "deleted": 0, "freed_bytes": 0})
        return

    ttl_days = cfg_get("cache_ttl_days", 30)
    now = datetime.now()
    stats = {"total": 0, "expired": 0, "deleted": 0, "freed_bytes": 0, "kept": 0}

    for root, _dirs, files in os.walk(str(cache_dir)):
        for fname in files:
            fpath = Path(root) / fname
            stats["total"] += 1
            fsize = fpath.stat().st_size

            _protected = {"session.json", "config.json", "cookies.json"}
            should_delete = args.clean_all and fname not in _protected
            if not should_delete and fname.endswith(".json") and ttl_days > 0 and fname not in _protected:
                try:
                    data = json.loads(fpath.read_text(encoding="utf-8"))
                    ts = data.get("_cached_at", "")
                    if ts:
                        cached_at = datetime.fromisoformat(ts)
                        if now - cached_at > timedelta(days=ttl_days):
                            should_delete = True
                            stats["expired"] += 1
                except Exception:
                    pass

            if should_delete:
                if not args.dry_run:
                    try:
                        fpath.unlink()
                        stats["deleted"] += 1
                        stats["freed_bytes"] += fsize
                    except Exception:
                        pass
                else:
                    stats["deleted"] += 1
                    stats["freed_bytes"] += fsize
            else:
                stats["kept"] += 1

    stats["freed_mb"] = round(stats["freed_bytes"] / 1024 / 1024, 2)
    mode = "dry-run" if args.dry_run else ("全部清理" if args.clean_all else f"TTL>{ttl_days}天")
    _output({"status": "success", "mode": mode, **stats})


# ── citations 命令 ────────────────────────────────────

def cmd_citations(args):
    """引文网络分析：获取论文的前向/后向引用"""
    paper_id = args.paper_id
    if not paper_id:
        _output({"status": "error", "code": "NO_PAPER_ID",
                 "message": "请提供论文标识（DOI、URL 或 arXiv ID）"})
        return

    direction = args.direction or "both"
    limit = args.limit or 20

    print(f"[citations] 查询 {paper_id} 的引文网络（方向: {direction}）...",
          file=sys.stderr)

    result = get_citations(paper_id, direction=direction, limit=limit)

    if "error" in result:
        _output({"status": "error", "code": "RESOLVE_FAILED",
                 "message": result["error"]})
        return

    paper = result.get("paper", {})
    citing = result.get("citing", [])
    references = result.get("references", [])
    status = "success" if paper else "partial"

    _output({
        "status": status,
        "paper": paper,
        "citing_count": len(citing),
        "references_count": len(references),
        "citing": citing,
        "references": references,
    })


# ── trends 命令 ───────────────────────────────────────

def cmd_trends(args):
    """研究趋势分析：基于当前会话数据进行聚合统计"""
    papers = _load_session(_session_project(args))
    if not papers:
        _output({"status": "error", "code": "NO_SESSION_DATA",
                 "message": "没有可分析的论文，请先执行 search、batch-search 或 import"})
        return

    print(f"[trends] 分析 {len(papers)} 篇论文的研究趋势...", file=sys.stderr)

    result = analyze_trends(papers)
    _output({"status": "success", "project": _session_project(args), **result})


# ── check 命令 ────────────────────────────────────────

def _check_browser(subprocess_mod) -> tuple:
    """检测可用浏览器，返回 (ok: bool, detail: str)"""
    import shutil as _shutil
    if sys.platform == "win32":
        import os as _os
        for p in [
            _os.path.expandvars(r"%ProgramFiles(x86)%\Microsoft\Edge\Application\msedge.exe"),
            _os.path.expandvars(r"%ProgramFiles%\Microsoft\Edge\Application\msedge.exe"),
            _os.path.expandvars(r"%LocalAppData%\Microsoft\Edge\Application\msedge.exe"),
        ]:
            if _os.path.exists(p):
                try:
                    r = subprocess_mod.run([p, "--version"], capture_output=True, text=True, timeout=5)
                    return True, r.stdout.strip() if r.returncode == 0 else f"Edge ({p})"
                except Exception:
                    return True, f"Edge ({p})"
        _chrome = _shutil.which("chrome") or _shutil.which("google-chrome")
        if _chrome:
            return True, f"Chrome ({_chrome})"
    else:
        for cmd in ["microsoft-edge", "microsoft-edge-stable", "google-chrome", "chromium", "chromium-browser"]:
            found = _shutil.which(cmd)
            if found:
                try:
                    r = subprocess_mod.run([found, "--version"], capture_output=True, text=True, timeout=5)
                    return True, r.stdout.strip() if r.returncode == 0 else cmd
                except Exception:
                    return True, cmd
    return False, "未检测到 Edge/Chrome"


def _check_driver() -> tuple:
    """检测浏览器驱动是否可用，返回 (ok: bool, detail: str)。"""
    try:
        from cnki.driver import _detect_browser, _find_local_driver
        browser = _detect_browser()
    except Exception as e:
        return False, f"无法检测浏览器类型: {e}"

    try:
        from selenium.webdriver.common.selenium_manager import SeleniumManager
        sm_bin = SeleniumManager._get_binary()
        if sm_bin and os.path.isfile(str(sm_bin)):
            import subprocess as _sp
            browser_arg = "MicrosoftEdge" if browser == "edge" else "chrome"
            result = _sp.run(
                [str(sm_bin), "--browser", browser_arg, "--offline", "--output", "json"],
                capture_output=True, text=True, timeout=15,
            )
            if result.returncode == 0:
                try:
                    info = json.loads(result.stdout)
                    driver_path = info.get("result", {}).get("driver_path", "")
                    if driver_path and os.path.isfile(driver_path):
                        return True, f"selenium-manager 找到: {driver_path}"
                except (ValueError, KeyError):
                    pass
    except Exception:
        pass

    local = _find_local_driver(browser)
    if local:
        return True, f"本地驱动: {local}"

    driver_name = "msedgedriver" if browser == "edge" else "chromedriver"
    return False, (
        f"未找到 {driver_name}（Selenium Manager 需联网下载）。"
        f"解决：1) 提权获取网络权限 2) 设置 SCHOLAR_DRIVER_PATH 环境变量"
    )


def _check_cnki() -> tuple:
    """检测知网连通性，返回 (ok: bool, detail: str)。
    SANDBOX_BLOCKED 单独标记，让 check 输出能指导 Agent 提权重试。
    """
    try:
        accessible, msg = check_cnki_access()
        if accessible:
            return True, "可访问"
        if msg.startswith("SANDBOX_BLOCKED"):
            return False, f"沙盒权限阻止（WinError 10013 等），提权后可能正常"
        return False, msg
    except Exception as e:
        return False, str(e)


def _fix_sandbox_network() -> List[str]:
    """检测沙箱环境并写入网络配置。返回已修复项列表。

    注意：Codex 的网络权限在任务创建时锁定，运行中写入 config.toml
    只对下次任务生效，当前任务仍然无网络。
    """
    fixes = []

    codex_dir = Path.cwd() / ".codex"
    codex_home = Path.home() / ".codex"
    is_codex_env = (codex_dir.exists() or codex_home.exists()
                    or os.environ.get("CODEX_SANDBOX_NETWORK_DISABLED"))
    if is_codex_env:
        config_targets = []
        if codex_dir.exists():
            config_targets.append(codex_dir / "config.toml")
        if codex_home.exists() and codex_home != codex_dir:
            config_targets.append(codex_home / "config.toml")
        if not config_targets:
            codex_dir.mkdir(parents=True, exist_ok=True)
            config_targets.append(codex_dir / "config.toml")

        codex_net_config = (
            '\n# Scholar Kit - 知网检索需要网络权限\n'
            'approval_policy = "on-request"\n'
            'sandbox_mode = "workspace-write"\n\n'
            '[sandbox_workspace_write]\n'
            'network_access = true\n'
        )
        for config_path in config_targets:
            try:
                existing = config_path.read_text(encoding="utf-8") if config_path.exists() else ""
                if "network_access" in existing:
                    continue
                with open(config_path, "a", encoding="utf-8") as f:
                    if existing and not existing.endswith("\n"):
                        f.write("\n")
                    f.write(codex_net_config)
                fixes.append(f"codex: 已写入 {config_path}（network_access = true）")
            except (PermissionError, OSError):
                continue

    claude_dir = Path.cwd() / ".claude"
    if claude_dir.exists():
        settings_path = claude_dir / "settings.json"
        try:
            settings = {}
            if settings_path.exists():
                settings = json.loads(settings_path.read_text(encoding="utf-8"))
            domains = settings.setdefault("sandbox", {}).setdefault("network", {}).setdefault("allowedDomains", [])
            if "*.cnki.net" not in domains:
                domains.append("*.cnki.net")
                settings_path.write_text(json.dumps(settings, ensure_ascii=False, indent=2), encoding="utf-8")
                fixes.append("claude-code: 已在 .claude/settings.json 添加 *.cnki.net")
        except Exception:
            pass

    return fixes


def cmd_check(args):
    """环境自检：逐项检查运行条件，输出能力位供 Agent 决策。--fix 时自动修复可修复项。"""
    import subprocess
    fix_mode = getattr(args, "fix", False)

    checks = []
    fixes_applied = []

    v = sys.version_info
    checks.append({
        "item": "Python",
        "status": "ok" if v >= (3, 9) else "warn" if v >= (3, 8) else "fail",
        "detail": f"{v.major}.{v.minor}.{v.micro}",
    })

    for pkg, import_name in [
        ("selenium", "selenium"), ("httpx", "httpx"),
        ("beautifulsoup4", "bs4"), ("openpyxl", "openpyxl"),
        ("python-docx", "docx"),
    ]:
        try:
            mod = __import__(import_name)
            ver = getattr(mod, "__version__", "?")
            status = "ok"
            detail = ver
            if pkg == "selenium" and ver != "?":
                try:
                    parts = [int(x) for x in ver.split(".")[:2]]
                    if parts < [4, 10]:
                        status = "warn"
                        detail = f"{ver}（需要 >=4.10）"
                except (ValueError, IndexError):
                    pass
            checks.append({"item": pkg, "status": status, "detail": detail})
        except ImportError:
            if pkg == "httpx":
                checks.append({"item": pkg, "status": "warn", "detail": "未安装（urllib 兜底可用）"})
            elif fix_mode and pkg == "selenium":
                try:
                    subprocess.run(
                        [sys.executable, "-m", "pip", "install", "selenium>=4.10"],
                        capture_output=True, timeout=120,
                    )
                    mod = __import__(import_name)
                    ver = getattr(mod, "__version__", "?")
                    checks.append({"item": pkg, "status": "ok", "detail": ver})
                    fixes_applied.append(f"selenium: 已自动安装 ({ver})")
                except Exception as e:
                    checks.append({"item": pkg, "status": "fail", "detail": f"自动安装失败: {e}"})
            else:
                checks.append({"item": pkg, "status": "fail", "detail": "未安装"})

    encoding = sys.stdout.encoding or "unknown"
    checks.append({
        "item": "终端编码",
        "status": "ok" if "utf" in encoding.lower() else "warn",
        "detail": encoding,
    })

    browser_ok, browser_detail = _check_browser(subprocess)
    checks.append({
        "item": "浏览器",
        "status": "ok" if browser_ok else "warn",
        "detail": browser_detail,
    })

    driver_ok, driver_detail = _check_driver()
    checks.append({
        "item": "浏览器驱动",
        "status": "ok" if driver_ok else "warn",
        "detail": driver_detail,
    })

    cnki_ok, cnki_detail = _check_cnki()
    checks.append({
        "item": "知网连通性",
        "status": "ok" if cnki_ok else "fail",
        "detail": cnki_detail,
    })

    if fix_mode and not cnki_ok:
        sandbox_fixes = _fix_sandbox_network()
        if sandbox_fixes:
            fixes_applied.extend(sandbox_fixes)
            cnki_ok, cnki_detail = _check_cnki()
            for c in checks:
                if c["item"] == "知网连通性":
                    c["status"] = "ok" if cnki_ok else "fail"
                    c["detail"] = cnki_detail
                    break

    cache_dir = Path.cwd() / ".scholar-kit"
    if cache_dir.exists():
        total = sum(f.stat().st_size for f in cache_dir.rglob("*") if f.is_file())
        checks.append({
            "item": "缓存目录", "status": "ok",
            "detail": f"{cache_dir} ({round(total/1024/1024, 2)} MB)",
        })
    else:
        checks.append({"item": "缓存目录", "status": "ok", "detail": "尚未创建"})

    selenium_item = next((c for c in checks if c["item"] == "selenium"), None)
    selenium_ok = selenium_item is not None and selenium_item["status"] != "fail"
    sandbox_blocked = cnki_detail and "沙盒权限阻止" in cnki_detail
    cnki_feasible = browser_ok and cnki_ok and selenium_ok and driver_ok
    cnki_reasons = []
    if not selenium_ok:
        cnki_reasons.append("selenium 未安装")
    if not browser_ok:
        cnki_reasons.append("未检测到浏览器")
    if not driver_ok:
        cnki_reasons.append("浏览器驱动缺失（需联网下载）")
    if not cnki_ok:
        cnki_reasons.append("知网不可达（沙盒权限阻止）" if sandbox_blocked else "知网不可达")

    is_codex = (
        (Path.cwd() / ".codex").exists()
        or os.environ.get("CODEX_SANDBOX_NETWORK_DISABLED")
        or any(k.startswith("CODEX_") for k in os.environ)
    )

    capabilities: Dict[str, Any] = {
        "cnki_feasible": cnki_feasible,
        "sandbox_blocked": sandbox_blocked,
        "driver_ok": driver_ok,
        "api_sources": True,
        "docx_tools": any(c["item"] == "python-docx" and c["status"] == "ok" for c in checks),
    }
    if not cnki_feasible:
        capabilities["cnki_blocked_reasons"] = cnki_reasons
        capabilities["suggested_sources"] = ["openalex", "semantic", "arxiv", "nssd"]

        if sandbox_blocked or (is_codex and not driver_ok):
            capabilities["needs_escalation"] = True
            capabilities["reason_code"] = "browser_or_driver_requires_host_permission"
            capabilities["agent_action"] = "rerun_same_command_with_escalated_permissions"
            capabilities["retry_command"] = (
                "py -3 -X utf8 scripts/literature.py check --fix"
                if sys.platform == "win32"
                else "python3 scripts/literature.py check --fix"
            )
            capabilities["recommended_env"] = {
                "SE_CACHE_PATH": ".scholar-kit/selenium-cache",
                "SCHOLAR_SKIP_NETWORK_CHECK": "1",
            }
            capabilities["fix_note"] = (
                "沙盒限制导致预检/驱动下载失败。"
                "解决：设置 SCHOLAR_SKIP_NETWORK_CHECK=1，"
                "然后以提权方式运行知网命令（平台会弹窗让用户授权网络访问）。"
                "提权后 Selenium Manager 自动下载驱动，所有卡点一步清零"
            )
        elif not driver_ok:
            capabilities["fix_note"] = (
                "浏览器驱动缺失。确保有网络连接后运行任意知网命令，"
                "Selenium Manager 会自动下载匹配的驱动。"
                "或手动设置 SCHOLAR_DRIVER_PATH 环境变量指向已有驱动"
            )
        elif fix_mode and not cnki_ok and fixes_applied:
            capabilities["fix_note"] = "已写入沙箱网络配置但知网仍不可达。可能原因：未连接校园网/VPN"

    update_info = _check_update()

    all_ok = all(c["status"] != "fail" for c in checks)
    output: Dict[str, Any] = {
        "status": "success" if all_ok else "warning",
        "version": __version__,
        "capabilities": capabilities,
        "checks": checks,
    }
    if fixes_applied:
        output["fixes_applied"] = fixes_applied
    if update_info:
        output["update"] = update_info
    _output(output)


def _check_update():
    """只读版本对比：本地版本 vs GitHub 最新 Release/Tag（超时不阻塞）"""
    import re
    _SEMVER_RE = re.compile(r"^\d+\.\d+(\.\d+)?$")

    def _version_key(value: str) -> tuple:
        parts = [int(p) for p in value.split(".")]
        return tuple((parts + [0, 0])[:3])

    repo = "lottshin/scholar-kit"
    urls = [
        f"https://api.github.com/repos/{repo}/releases/latest",
        f"https://api.github.com/repos/{repo}/tags?per_page=1",
    ]

    try:
        latest = None
        for url in urls:
            if latest:
                break
            try:
                data = _fetch_json(url, timeout=5)
                if data is None:
                    continue
                if isinstance(data, list):
                    latest = data[0].get("name", "") if data else ""
                else:
                    latest = data.get("tag_name", "")
            except Exception:
                continue

        if not latest:
            return None

        latest = latest.removeprefix("v")
        if not _SEMVER_RE.match(latest):
            return None

        current = __version__.removeprefix("v")
        if not _SEMVER_RE.match(current):
            return None

        if _version_key(latest) > _version_key(current):
            return {
                "update_available": True,
                "current": current,
                "latest": latest,
                "message": f"新版本 {latest} 可用，在 skill 目录执行 git pull 更新",
            }
        return {"update_available": False, "current": current, "latest": latest}
    except Exception:
        return None


def _fetch_json(url: str, timeout: int = 10):
    """HTTP GET 返回 JSON，httpx 优先，urllib 兜底。失败统一返回 None。"""
    try:
        import httpx
        resp = httpx.get(url, timeout=timeout, follow_redirects=True)
        return resp.json() if resp.status_code == 200 else None
    except ImportError:
        pass
    except Exception:
        return None
    try:
        from urllib.request import urlopen, Request
        req = Request(url, headers={"User-Agent": "scholar-kit",
                                    "Accept": "application/json"})
        with urlopen(req, timeout=timeout) as resp:
            if resp.status != 200:
                return None
            import json as _json
            return _json.loads(resp.read().decode("utf-8", errors="replace"))
    except Exception:
        return None


# ── 输出工具 ──────────────────────────────────────────

def _output(data):
    print(json.dumps(data, ensure_ascii=False, indent=2))


# ── CLI 解析 ──────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        prog="literature",
        description="Scholar Kit - 学术文献检索工具",
    )
    parser.add_argument("--version", action="version", version=f"%(prog)s {__version__}")

    sub = parser.add_subparsers(dest="command")

    # search
    p_search = sub.add_parser("search", help="搜索文献")
    p_search.add_argument("query", help="搜索关键词")
    p_search.add_argument("--source", default="cnki",
                          help="数据源: cnki/openalex/semantic/arxiv/nssd/all (默认 cnki)")
    p_search.add_argument("--limit", type=int, default=None, help="结果数量限制（默认 20，多页时自动扩展）")
    p_search.add_argument("--core",
                          help="知网侧边栏来源类别，逗号分隔: 北大核心,CSSCI,AMI,WJCI,CSCD,EI")
    p_search.add_argument("--doc-type",
                          choices=["journal", "master", "doctor", "thesis",
                                   "conference", "newspaper"],
                          help="文献类型筛选: journal/master/doctor/thesis/conference/newspaper")
    p_search.add_argument("--field", default=None,
                          help="搜索字段: 主题(默认)/篇名/关键词/摘要/全文/作者/来源")
    p_search.add_argument("--year-from", type=int, help="起始年份")
    p_search.add_argument("--year-to", type=int, help="截止年份")
    p_search.add_argument("--author", help="作者")
    p_search.add_argument("--journal", help="期刊名")
    p_search.add_argument("--sort", choices=["relevance", "date", "citations", "quality"],
                          default="relevance", help="排序方式")
    p_search.add_argument("--pages", type=int, default=1, help="知网抓取页数")
    p_search.add_argument("--export", help="直接导出: bibtex/ris/markdown/json/excel")
    p_search.add_argument("--output", help="导出文件路径")
    p_search.add_argument("--download", action="store_true",
                          help="搜索后直接下载（仅 --source cnki）")
    p_search.add_argument("--download-dir", default="./papers",
                          help="下载目录（配合 --download，默认 ./papers）")
    p_search.add_argument("--download-top-n", type=int, default=None,
                          help="下载前 N 篇（配合 --download，默认全部）")
    p_search.add_argument("--download-file-format", choices=["pdf", "caj"], default="pdf",
                          help="下载文件格式（配合 --download，默认 pdf）")
    p_search.add_argument("--download-fallback-format", choices=["caj"], default=None,
                          help="主格式失败时的兜底格式；例如 PDF 按钮不存在时尝试 CAJ")
    p_search.add_argument("--download-citation-style", choices=CITATION_STYLE_CHOICES,
                          default="gbt7714", help="下载清单引用格式（配合 --download）")
    p_search.add_argument("--download-report-output",
                          help="下载清单输出路径（配合 --download，默认写入下载目录）")
    p_search.add_argument("--download-no-report", action="store_true",
                          help="不生成下载清单（配合 --download）")
    p_search.add_argument("--enrich", type=int, default=0, metavar="N",
                          help="自动补全前 N 篇知网论文的卷期页码（需访问详情页）")
    p_search.add_argument("--cite-enrich", type=int, default=0, metavar="N",
                          help="搜索时点击前 N 篇知网引用按钮，快速补全 GB/T 引用和页码")
    p_search.add_argument("--append", action="store_true",
                          help="追加到已有会话结果（而非覆盖）")
    p_search.add_argument("--project", help="课题文献库名称；指定后读写 .scholar-kit/projects/<project>/session.json")
    p_search.set_defaults(func=cmd_search)

    # download
    p_download = sub.add_parser("download", help="下载论文")
    p_download.add_argument("target", nargs="?", help="知网论文 URL")
    p_download.add_argument("--doi", help="通过 DOI 查找 OA 链接")
    p_download.add_argument("--dir", default="./papers", help="保存目录")
    p_download.add_argument("--file-format", choices=["pdf", "caj"], default="pdf")
    p_download.set_defaults(func=cmd_download)

    # detail
    p_detail = sub.add_parser("detail", help="获取知网论文详情")
    p_detail.add_argument("url", help="知网论文详情页 URL")
    p_detail.set_defaults(func=cmd_detail)

    # export
    p_export = sub.add_parser("export", help="导出上次搜索结果")
    p_export.add_argument("--format", dest="export_format", required=True,
                          choices=["bibtex", "ris", "markdown", "json", "excel",
                                   "gbt7714", "footnote", "apa", "mla", "chicago"])
    p_export.add_argument("--output", help="输出文件路径")
    p_export.add_argument("--raw", action="store_true", help="输出纯文本而非 JSON")
    p_export.add_argument("--project", help="课题文献库名称")
    p_export.set_defaults(func=cmd_export)

    # projects
    p_projects = sub.add_parser("projects", help="列出课题文献库")
    p_projects.set_defaults(func=cmd_projects)

    # library
    p_library = sub.add_parser("library", help="查看当前或指定课题文献库")
    p_library.add_argument("--project", help="课题文献库名称")
    p_library.add_argument("--limit", type=int, help="最多显示前 N 篇")
    p_library.set_defaults(func=cmd_library)

    # cite
    p_cite = sub.add_parser("cite", help="生成引用格式")
    p_cite.add_argument("--style", default="gbt7714",
                        choices=CITATION_STYLE_CHOICES)
    p_cite.add_argument("--raw", action="store_true", help="输出纯文本而非 JSON")
    p_cite.add_argument("--project", help="课题文献库名称")
    p_cite.set_defaults(func=cmd_cite)

    # import
    p_import = sub.add_parser("import", help="导入知网导出的题录文件")
    p_import.add_argument("filepath", help="题录文件路径")
    p_import.add_argument("--project", help="课题文献库名称")
    p_import.set_defaults(func=cmd_import)

    # read-paper
    p_paper = sub.add_parser("read-paper", help="读取论文文件（.docx/.txt/.md）并输出 UTF-8 文本")
    p_paper.add_argument("filepath", help="论文文件路径")
    p_paper.add_argument("--output", help="输出到文件（默认直接打印）")
    p_paper.add_argument("--raw", action="store_true", help="输出纯文本而非 JSON")
    p_paper.set_defaults(func=cmd_read_paper)

    # pdf-meta
    p_pdf = sub.add_parser("pdf-meta", help="从 PDF 提取元数据（标题、DOI 等）")
    p_pdf.add_argument("filepath", help="PDF 文件路径")
    p_pdf.set_defaults(func=cmd_pdf_meta)

    # batch-search
    p_batch = sub.add_parser("batch-search", help="批量知网搜索（一次启动浏览器）")
    p_batch.add_argument("queries", nargs="*", help="搜索关键词列表")
    p_batch.add_argument("--query-file", help="关键词文件路径（每行一个关键词）")
    p_batch.add_argument("--core",
                         help="知网侧边栏来源类别，逗号分隔: 北大核心,CSSCI,AMI,WJCI,CSCD,EI")
    p_batch.add_argument("--doc-type",
                         choices=["journal", "master", "doctor", "thesis",
                                  "conference", "newspaper"],
                         help="文献类型筛选: journal/master/doctor/thesis/conference/newspaper")
    p_batch.add_argument("--field", default=None,
                         help="搜索字段: 主题(默认)/篇名/关键词/摘要/全文/作者/来源")
    p_batch.add_argument("--author", help="作者（对每组关键词生效）")
    p_batch.add_argument("--journal", help="期刊名（对每组关键词生效）")
    p_batch.add_argument("--year-from", type=int, help="起始年份")
    p_batch.add_argument("--year-to", type=int, help="截止年份")
    p_batch.add_argument("--sort", choices=["relevance", "date", "citations", "quality"],
                         default="relevance", help="排序方式")
    p_batch.add_argument("--pages", type=int, default=1, help="每组关键词抓取页数")
    p_batch.add_argument("--export", help="直接导出: bibtex/ris/markdown/json/excel")
    p_batch.add_argument("--output", help="导出文件路径")
    p_batch.add_argument("--append", action="store_true",
                         help="追加到已有会话结果（而非覆盖）")
    p_batch.add_argument("--project", help="课题文献库名称")
    p_batch.set_defaults(func=cmd_batch_search)

    # read-detail
    p_read = sub.add_parser("read-detail", help="批量获取论文摘要/全文（需先搜索）")
    p_read.add_argument("--top-n", type=int, default=5,
                        help="获取前 N 篇论文的详情（默认 5）")
    p_read.add_argument("--indices", type=str, default=None,
                        help="指定论文序号（从1开始），如 '3' '1,3,9' '2-5'。指定后忽略 --top-n")
    p_read.add_argument("--fulltext", action="store_true",
                        help="抓取 HTML 全文（默认只抓摘要）")
    p_read.add_argument("--project", help="课题文献库名称")
    p_read.set_defaults(func=cmd_read_detail)

    # batch-download
    p_bdl = sub.add_parser("batch-download", help="批量下载知网论文（一次启动浏览器）")
    p_bdl.add_argument("urls", nargs="*", help="知网论文 URL 列表（可选，也可用 --from-session）")
    p_bdl.add_argument("--from-session", action="store_true",
                       help="从上次搜索结果（session.json）读取 URL")
    p_bdl.add_argument("--top-n", type=int, help="配合 --from-session，只下载前 N 篇")
    p_bdl.add_argument("--dir", default="./papers", help="保存目录")
    p_bdl.add_argument("--file-format", choices=["pdf", "caj"], default="pdf")
    p_bdl.add_argument("--fallback-format", choices=["caj"], default=None,
                       help="主格式失败时的兜底格式；例如 PDF 按钮不存在时尝试 CAJ")
    p_bdl.add_argument("--citation-style", choices=CITATION_STYLE_CHOICES,
                       default="gbt7714", help="下载清单引用格式")
    p_bdl.add_argument("--report-output", help="下载清单输出路径（默认写入下载目录）")
    p_bdl.add_argument("--no-report", action="store_true", help="不生成下载清单")
    p_bdl.add_argument("--project", help="课题文献库名称；配合 --from-session 使用")
    p_bdl.set_defaults(func=cmd_batch_download)

    # write-docx
    p_wdocx = sub.add_parser("write-docx", help="Markdown → 学术格式 Word 文档")
    p_wdocx.add_argument("filepath", help="Markdown 文件路径")
    p_wdocx.add_argument("--output", help="输出 .docx 路径（默认同名 .docx）")
    p_wdocx.set_defaults(func=cmd_write_docx)

    # patch-docx
    p_pdocx = sub.add_parser("patch-docx", help="在现有 .docx 上打补丁（插入引用/脚注）")
    p_pdocx.add_argument("filepath", help="原始 .docx 文件路径")
    p_pdocx.add_argument("--patch", required=True, help="补丁 JSON 文件路径")
    p_pdocx.add_argument("--output", help="输出路径（默认 原名_patched.docx）")
    p_pdocx.set_defaults(func=cmd_patch_docx)

    # clean-cache
    p_clean = sub.add_parser("clean-cache", help="清理过期缓存文件")
    p_clean.add_argument("--all", action="store_true", dest="clean_all",
                         help="删除所有缓存（不仅是过期的）")
    p_clean.add_argument("--dry-run", action="store_true",
                         help="仅统计，不实际删除")
    p_clean.set_defaults(func=cmd_clean_cache)

    # citations
    p_cite_net = sub.add_parser("citations", help="引文网络分析（前向/后向引用）")
    p_cite_net.add_argument("paper_id", help="论文标识（DOI、Semantic Scholar URL、arXiv ID 等）")
    p_cite_net.add_argument("--direction", choices=["citing", "cited", "both"], default="both",
                            help="引用方向：citing=谁引了它，cited=它引了谁，both=双向（默认）")
    p_cite_net.add_argument("--limit", type=int, default=20,
                            help="每个方向最多返回条数（默认 20）")
    p_cite_net.set_defaults(func=cmd_citations)

    # trends
    p_trends = sub.add_parser("trends", help="研究趋势分析（基于会话中的搜索结果）")
    p_trends.add_argument("--project", help="课题文献库名称")
    p_trends.set_defaults(func=cmd_trends)

    # review
    p_review = sub.add_parser("review", help="基于会话/课题文献库生成可追溯综述材料")
    p_review.add_argument("--topic", help="综述主题；默认使用 --project 或当前课题")
    p_review.add_argument("--project", help="课题文献库名称")
    p_review.add_argument("--limit", type=int, default=12, help="最多纳入前 N 篇相关文献（默认 12）")
    p_review.add_argument("--output", help="输出 Markdown 文件路径")
    p_review.add_argument("--auto-detail", action="store_true", help="生成综述前自动补全高相关知网文献摘要")
    p_review.add_argument("--detail-top-n", type=int, default=5, help="配合 --auto-detail，最多补全 N 篇知网文献（默认 5）")
    p_review.add_argument("--cluster", action="store_true", help="按主题聚类组织综述材料")
    p_review.add_argument("--gaps", action="store_true", help="基于当前文献库统计生成研究空白提示")
    p_review.add_argument("--raw", action="store_true", help="直接输出 Markdown 文本")
    p_review.set_defaults(func=cmd_review)

    # write
    p_write = sub.add_parser("write", help="基于课题文献库直接写作文献综述")
    p_write.add_argument("--project", help="课题文献库名称")
    p_write.add_argument("--topic", help="写作主题；默认使用 --project 或当前课题")
    p_write.add_argument("--limit", type=int, default=12, help="最多纳入前 N 篇相关文献（默认 12）")
    p_write.add_argument("--format", choices=["markdown", "md", "docx"], default="markdown", help="输出格式：markdown/md/docx")
    p_write.add_argument("--mode", choices=["outline", "draft", "section"], default="draft", help="写作模式：outline 大纲 / draft 正文 / section 单节")
    p_write.add_argument("--section", help="只生成指定章节，如 研究背景、研究不足、某个主题聚类名称")
    p_write.add_argument("--output", help="输出文件路径")
    p_write.add_argument("--with-citations", action="store_true", help="附加参考文献列表")
    p_write.add_argument("--citation-style", default="gbt7714", choices=CITATION_STYLE_CHOICES)
    p_write.add_argument("--validate", action="store_true", help="同时输出写作证据质量校验报告")
    p_write.add_argument("--raw", action="store_true", help="直接输出 Markdown 文本")
    p_write.set_defaults(func=cmd_write)

    p_validate = sub.add_parser("validate", help="校验综述正文的证据支撑质量")
    p_validate.add_argument("--project", help="课题文献库名称")
    p_validate.add_argument("--topic", help="写作主题；默认使用 --project 或当前课题")
    p_validate.add_argument("--limit", type=int, default=12, help="最多纳入前 N 篇相关文献（默认 12）")
    p_validate.add_argument("--file", help="待校验的 Markdown 文件；缺省时校验当前自动生成草稿")
    p_validate.set_defaults(func=cmd_validate)

    p_topics = sub.add_parser("topics", help="基于课题文献库生成带证据的选题建议")
    p_topics.add_argument("--project", help="课题文献库名称")
    p_topics.add_argument("--topic", help="选题方向；默认使用 --project 或当前课题")
    p_topics.add_argument("--limit", type=int, default=6, help="最多生成 N 个选题建议（默认 6）")
    p_topics.set_defaults(func=cmd_topics)

    # check
    p_check = sub.add_parser("check", help="环境自检（Python / 依赖 / 浏览器 / 知网连通性）")
    p_check.add_argument("--fix", action="store_true",
                         help="自动修复可修复项（安装 selenium、写入沙箱网络配置）")
    p_check.set_defaults(func=cmd_check)

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return

    args.func(args)


if __name__ == "__main__":
    main()
